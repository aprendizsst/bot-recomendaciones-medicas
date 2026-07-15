import streamlit as st
import pdfplumber
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
import sqlite3
import hashlib
import os
import datetime
import re
import unicodedata
try:
    import pytesseract
    from PIL import ImageEnhance, ImageFilter, ImageOps
    OCR_DISPONIBLE = True
except Exception:
    pytesseract = None
    ImageEnhance = None
    ImageFilter = None
    ImageOps = None
    OCR_DISPONIBLE = False
import requests
import io
import zipfile
import tempfile
import subprocess
import base64

# --- CONFIGURACIÓN DE PÁGINA AVANZADA ---
st.set_page_config(
    page_title="Portal SST - JER S.A.", 
    page_icon="🩺", 
    layout="wide"
)

# --- INYECCIÓN DE CSS: MODO OSCURO PREMIUM ---
st.markdown("""
    <style>
    html, body, [data-testid="stAppViewContainer"], [data-testid="stSidebar"], [data-testid="stHeader"] {
        background-color: #0b0f19 !important;
        color: #f8fafc !important;
    }
    
    .login-box {
        max-width: 450px;
        margin: 80px auto;
        padding: 40px;
        background-color: #111827 !important;
        border-radius: 16px;
        box-shadow: 0 20px 40px rgba(0, 0, 0, 0.4);
        border: 1px solid #1f2937;
    }
    
    .login-box h2 {
        color: #3b82f6 !important;
        text-align: center;
        margin-bottom: 5px;
        font-weight: 700;
    }
    .login-box p {
        color: #9ca3af !important;
        text-align: center;
        font-size: 14px;
    }
    
    div[data-testid="stRadio"] label p {
        color: #f3f4f6 !important;
        font-weight: 600 !important;
        font-size: 14px !important;
    }
    
    div[data-testid="stWidgetLabel"] p {
        color: #60a5fa !important;
        font-weight: 600 !important;
        font-size: 15px !important;
    }
    
    div[data-baseweb="input"] {
        background-color: #1f2937 !important;
        border: 1px solid #374151 !important;
        border-radius: 8px !important;
    }
    div[data-baseweb="input"] input {
        color: #ffffff !important;
        background-color: #1f2937 !important;
    }
    div[data-testid="stTextArea"] textarea {
        color: #ffffff !important;
        background-color: #1f2937 !important;
        border: 1px solid #374151 !important;
        border-radius: 8px !important;
    }
    
    button[data-baseweb="tab"] p {
        color: #9ca3af !important;
    }
    button[data-baseweb="tab"][aria-selected="true"] p {
        color: #3b82f6 !important;
        font-weight: 700 !important;
    }
    
    .stButton>button {
        background: linear-gradient(135deg, #2563eb 0%, #1d4ed8 100%) !important;
        color: #ffffff !important;
        border-radius: 8px !important;
        padding: 12px 24px !important;
        font-weight: 700 !important;
        border: none !important;
        transition: all 0.2s ease-in-out !important;
        box-shadow: 0 4px 14px rgba(37, 99, 235, 0.3) !important;
        width: 100%;
    }
    .stButton>button:hover {
        transform: translateY(-1px) !important;
        box-shadow: 0 6px 20px rgba(37, 99, 235, 0.5) !important;
        filter: brightness(115%);
    }
    
    .header-banner {
        background: linear-gradient(135deg, #1e3a8a 0%, #0f172a 100%);
        padding: 22px;
        border-radius: 12px;
        color: #ffffff !important;
        margin-bottom: 25px;
        border: 1px solid #1e2937;
        box-shadow: 0 4px 15px rgba(0, 0, 0, 0.3);
    }
    
    .metric-card {
        background-color: #111827 !important;
        padding: 15px;
        border-radius: 10px;
        box-shadow: 0 4px 10px rgba(0,0,0,0.2);
        border-left: 5px solid #2563eb;
        margin-bottom: 12px;
        color: #e5e7eb !important;
    }
    
    div[data-testid="stExpander"] {
        background-color: #111827 !important;
        border: 1px solid #1f2937 !important;
        border-radius: 8px !important;
    }
    </style>
""", unsafe_allow_html=True)

# --- SISTEMA DE SEGURIDAD (BASE DE DATOS) ---
DB_NAME = "usuarios.db"

def init_db():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS usuarios 
                 (usuario TEXT PRIMARY KEY, contrasena TEXT, nombre TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS configuracion 
                 (clave TEXT PRIMARY KEY, valor TEXT)''')
    conn.commit()
    conn.close()

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

def registrar_usuario(user, pwd, name):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    try:
        c.execute("INSERT INTO usuarios VALUES (?, ?, ?)", (user.lower().strip(), hash_password(pwd), name.strip()))
        conn.commit()
        success = True
    except sqlite3.IntegrityError:
        success = False
    conn.close()
    return success

def verificar_usuario(user, pwd):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT nombre FROM usuarios WHERE usuario = ? AND contrasena = ?", (user.lower().strip(), hash_password(pwd)))
    resultado = c.fetchone()
    conn.close()
    return resultado[0] if resultado else None

def actualizar_contrasena(user, old_pwd, new_pwd):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT contrasena FROM usuarios WHERE usuario = ?", (user.lower().strip(),))
    resultado = c.fetchone()
    if resultado and resultado[0] == hash_password(old_pwd):
        c.execute("UPDATE usuarios SET contrasena = ? WHERE usuario = ?", (hash_password(new_pwd), user.lower().strip()))
        conn.commit()
        conn.close()
        return True
    conn.close()
    return False

def tiene_usuarios():
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT COUNT(*) FROM usuarios")
    count = c.fetchone()[0]
    conn.close()
    return count > 0

def guardar_config(clave, valor):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("INSERT OR REPLACE INTO configuracion VALUES (?, ?)", (clave, valor))
    conn.commit()
    conn.close()

def obtener_config(clave):
    conn = sqlite3.connect(DB_NAME)
    c = conn.cursor()
    c.execute("SELECT valor FROM configuracion WHERE clave = ?", (clave,))
    res = c.fetchone()
    conn.close()
    return res[0] if res else ""

init_db()

# --- MANEJO DE SESIÓN DE LOGIN Y ESTADOS DE ARCHIVOS ---
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
if "username" not in st.session_state:
    st.session_state.username = ""
if "documentos" not in st.session_state:
    st.session_state.documentos = {}
if "pdfs_raw_bytes" not in st.session_state:
    st.session_state.pdfs_raw_bytes = {}
if "textos_raw" not in st.session_state:
    st.session_state.textos_raw = {}
if "export_bytes" not in st.session_state:
    st.session_state.export_bytes = None
if "zip_bytes" not in st.session_state:
    st.session_state.zip_bytes = None
if "processed_doc" not in st.session_state:
    st.session_state.processed_doc = None
if "prev_colaborador" not in st.session_state:
    st.session_state.prev_colaborador = None
if "document_count" not in st.session_state:
    st.session_state.document_count = 0

# --- REVISOR Y CORRECTOR DE ORTOGRAFÍA SST ---
def corregir_ortografia_sst(texto):
    if not texto: return ""
    diccionario_SST = {
        r'\brealziado\b': 'realizado', r'\brealziados\b': 'realizados',
        r'\baudiometria\b': 'audiometría', r'\bvisiometria\b': 'visiometría',
        r'\bespirometria\b': 'espirometría', r'\boptometria\b': 'optometría',
        r'\bfisica\b': 'física', r'\bmedico\b': 'médico', r'\bperiodico\b': 'periódico',
        r'\bproteccion\b': 'protección', r'\balimentacion\b': 'alimentación',
        r'\brecomendacion\b': 'recomendación', r'\bperfil\s+lipidico\b': 'perfil lipídico',
        r'\benfasis\b': 'énfasis', r'\bosteomuscular\b': 'osteomuscular',
        r'\bregion\b': 'región', r'\bhabitos\b': 'hábitos', r'\badiministrativo\b': 'administrativo'
    }
    for patron, reemplazo in diccionario_SST.items():
        texto = re.sub(patron, reemplazo, texto, flags=re.IGNORECASE)
    return texto

def a_caso_oracion(texto):
    if not texto: return ""
    texto_min = corregir_ortografia_sst(texto).lower().strip()
    return re.sub(r'(^|[.!?]\s+|\n+)([a-zñáéíóúü])', lambda m: m.group(1) + m.group(2).upper(), texto_min)

def es_vacio_o_negativo(texto):
    if not texto: return True
    return texto.strip().lower().strip(" .-_/ '\"") in ["no", "ninguna", "ninguno", "no registra", "sin remisiones", "normal", "n/a", "sin remisión"]

def es_vacio_o_estado(texto):
    if not texto: return True
    t_clean = texto.strip().upper()
    t_clean_norm = re.sub(r'[^A-ZÁÉÍÓÚÑ\s]', '', t_clean).strip()
    t_clean_norm = re.sub(r'\s+', ' ', t_clean_norm)
    
    frases_estado = {
        "REALIZADO", "REALZIADO", "SIN ALTERACIONES", "NORMAL", "SANO", "NEGATIVO", 
        "NO REGISTRA", "NA", "SIN REMISIONES", "SIN REMISIÓN", "VISUAL", "CARDIOVASCULAR", 
        "DME", "OSTEOMUSCULAR", "AUDITIVO", "RESPIRATORIO", "SVE", "SISTEMA", "VIGILANCIA",
        "SANO Y SIN ALTERACIONES", "NINGUNO", "NINGUNA", "NO PRESENTAS", "NO PRESENTA", 
        "NO REGISTRA RECOMENDACIONES", "NORMALES", "NORMAL", "SIN ALTERACION", "NO APLICA",
        "RECOMENDACIONES MÉDICAS", "RECOMENDACIONES OCUPACIONALES", "HABITOS Y ESTILO DE VIDA SALUDABLES",
        "HABITOS SALUDABLES", "OTRAS OBSERVACIONES Y RECOMENDACIONES", "RECOMENDACIONES MEDICAS"
    }
    
    if t_clean_norm in frases_estado:
        return True
    if len(texto.strip()) <= 3:
        return True
    return False

def limpiar_campo(texto):
    if not texto: return ""
    partes = re.split(r'\b(Teléfono|Telefono|Tel|C\.C|CC|Documento|Cedula|Cargo|Fecha)\b', texto, flags=re.IGNORECASE)
    return re.sub(r'[:\-,_]+', '', partes[0]).strip().title()

def limpiar_linea_ruido_lateral(linea):
    patron_ruido = r'\s{2,}(VISUAL|DME|CARDIOVASCULAR|SVE|AUDITIVO|RESPIRATORIO|SISTEMA|VIGILANCIA)\s*$'
    linea_limpia = re.sub(patron_ruido, '', linea, flags=re.IGNORECASE)
    return linea_limpia.strip()

def limpiar_ruido_columnas_final(texto):
    if not texto: return ""
    patrones_ruido = [
        r'\bvisual\b', r'\bdme\b', r'\bcardiovascular\b', r'\bsve\b', 
        r'\bauditivo\b', r'\brespiratorio\b', r'\bsistema\b', r'\bvigilancia\b'
    ]
    for patron in patrones_ruido:
        texto = re.sub(patron + r'\s*$', '', texto, flags=re.IGNORECASE)
    return texto.strip(" :-,_/")

def intentar_parsear_fecha(fecha_str):
    fecha_str = fecha_str.lower().strip(" :-,_/.()[]|")
    m_letras = re.search(r'(\d{1,2})\s+de\s+([a-zñáéíóúü]+)\s+de\s+(20\d{2})', fecha_str)
    if m_letras:
        dia = int(m_letras.group(1))
        mes_str = m_letras.group(2)
        anio = int(m_letras.group(3))
        meses = {
            "enero": 1, "febrero": 2, "marzo": 3, "abril": 4, "mayo": 5, "junio": 6,
            "julio": 7, "agosto": 8, "septiembre": 9, "octubre": 10, "noviembre": 11, "diciembre": 12
        }
        return datetime.date(anio, meses.get(mes_str, 1), dia)
        
    m_ymd = re.search(r'(20\d{2})[-/](\d{1,2})[-/](\d{1,2})', fecha_str)
    if m_ymd: return datetime.date(int(m_ymd.group(1)), int(m_ymd.group(2)), int(m_ymd.group(3)))
        
    m_dmy = re.search(r'(\d{1,2})[-/](\d{1,2})[-/](20\d{2})', fecha_str)
    if m_dmy: return datetime.date(int(m_dmy.group(3)), int(m_dmy.group(2)), int(m_dmy.group(1)))
        
    return datetime.date.today()


# --- EXTRACCIÓN ROBUSTA DE IDENTIDAD, CARGO Y LUGAR ---
# Estas funciones complementan el analizador existente sin cambiar su flujo,
# sus campos de salida ni la generación de documentos.

_ETIQUETAS_CORTE = [
    "APELLIDOS Y NOMBRES", "NOMBRES Y APELLIDOS", "NOMBRE DEL TRABAJADOR",
    "NOMBRE TRABAJADOR", "NOMBRE COMPLETO", "TRABAJADOR", "PACIENTE",
    "CARGO ACTUAL", "CARGO", "OCUPACIÓN", "OCUPACION", "OFICIO", "PUESTO",
    "DOCUMENTO", "IDENTIFICACIÓN", "IDENTIFICACION", "CÉDULA", "CEDULA", "C.C.",
    "CC", "GÉNERO", "GENERO", "EDAD", "TELÉFONO", "TELEFONO", "CELULAR",
    "EPS", "AFP", "ARL", "EMPRESA", "NIT", "FECHA", "CIUDAD", "MUNICIPIO",
    "LUGAR", "SEDE", "DIRECCIÓN", "DIRECCION"
]

_RUIDO_IDENTIDAD = {
    "DATOS", "DATOS DEL TRABAJADOR", "INFORMACION DEL TRABAJADOR",
    "INFORMACIÓN DEL TRABAJADOR", "APELLIDOS Y NOMBRES", "NOMBRES Y APELLIDOS",
    "NOMBRE", "TRABAJADOR", "PACIENTE", "GENERO", "GÉNERO", "EDAD",
    "DOCUMENTO", "IDENTIFICACION", "IDENTIFICACIÓN", "CEDULA", "CÉDULA",
    "EMPRESA", "IPS", "EPS", "AFP", "ARL", "FIRMA", "CERTIFICADO"
}

_RUIDO_CARGO = {
    "CARGO", "CARGO ACTUAL", "OCUPACION", "OCUPACIÓN", "OFICIO", "PUESTO",
    "TRABAJADOR", "DATOS", "EMPRESA", "EPS", "AFP", "ARL", "GENERO", "GÉNERO",
    "DOCUMENTO", "IDENTIFICACION", "IDENTIFICACIÓN", "CERTIFICADO"
}

_RUIDO_LUGAR = {
    "LUGAR", "CIUDAD", "MUNICIPIO", "SEDE", "FECHA", "DIA", "DÍA", "MES",
    "AÑO", "ANO", "REALIZACION", "REALIZACIÓN", "EXAMEN", "EXÁMEN",
    "CERTIFICADO", "PAGINA", "PÁGINA", "LOGOTIPO", "AM", "PM"
}


def normalizar_etiqueta(texto):
    """Normaliza únicamente para comparar etiquetas; no modifica el valor final."""
    if not texto:
        return ""
    texto = unicodedata.normalize("NFKD", str(texto))
    texto = "".join(c for c in texto if not unicodedata.combining(c))
    texto = texto.upper()
    texto = re.sub(r"\s+", " ", texto)
    return texto.strip(" :|_-./")


def dividir_columnas_estructuradas(linea):
    """Divide filas de tablas extraídas por pdfplumber sin romper nombres normales."""
    if not linea:
        return []
    columnas = [
        re.sub(r"\s+", " ", c).strip(" |/-,_.:")
        for c in re.split(r"\s{2,}|\t+|\|", linea)
    ]
    return [c for c in columnas if c]


def recortar_en_siguiente_etiqueta(valor):
    if not valor:
        return ""
    patron = r"\b(?:" + "|".join(
        sorted((re.escape(e) for e in _ETIQUETAS_CORTE), key=len, reverse=True)
    ) + r")\b"
    coincidencias = list(re.finditer(patron, valor, flags=re.IGNORECASE))
    if coincidencias:
        # Se conserva la primera parte solo cuando la siguiente etiqueta aparece
        # después de haber extraído algún valor.
        primera = coincidencias[0]
        if primera.start() > 0:
            valor = valor[:primera.start()]
    return valor.strip(" |/-,_.:")


def limpiar_candidato_campo(valor, tipo):
    if not valor:
        return ""
    valor = str(valor).replace("\x00", " ")
    valor = re.sub(r"\s+", " ", valor).strip(" |/-,_.:")
    valor = recortar_en_siguiente_etiqueta(valor)

    # Elimina identificaciones y datos laterales que suelen quedar pegados.
    valor = re.split(
        r"\b(?:C\.?\s*C\.?|CÉDULA|CEDULA|DOCUMENTO|IDENTIFICACIÓN|IDENTIFICACION|"
        r"TELÉFONO|TELEFONO|CELULAR|EDAD|GÉNERO|GENERO|EPS|AFP|ARL)\b",
        valor,
        maxsplit=1,
        flags=re.IGNORECASE,
    )[0].strip(" |/-,_.:")

    if tipo in {"nombre", "cargo"}:
        valor = re.sub(r"\([^)]*\)", "", valor).strip()
    elif tipo == "lugar":
        # Conserva nombres de IPS entre paréntesis solo si son el único dato.
        valor = re.sub(r"\s+\([^)]*(?:PÁGINA|PAGINA|HORA|AM|PM)[^)]*\)", "", valor, flags=re.IGNORECASE)
        # En algunas tablas la línea siguiente contiene DÍA | MES | AÑO | CIUDAD.
        # Se eliminan los componentes de fecha antes de validar el municipio.
        valor = re.sub(
            r"^\s*(?:(?:20\d{2})\s*[|/\-.]\s*\d{1,2}\s*[|/\-.]\s*\d{1,2}|"
            r"\d{1,2}\s*[|/\-.]\s*\d{1,2}\s*[|/\-.]\s*20\d{2})\s*[|/\-,_.:]*\s*",
            "",
            valor,
        )
        if "|" in valor:
            partes_lugar = [
                p.strip(" |/-,_.:")
                for p in valor.split("|")
                if p.strip(" |/-,_.:")
            ]
            partes_con_letras = [
                p for p in partes_lugar
                if sum(c.isalpha() for c in p) >= 3
            ]
            if partes_con_letras:
                valor = partes_con_letras[-1]

    return re.sub(r"\s+", " ", valor).strip(" |/-,_.:")


def candidato_nombre_valido(valor):
    valor = limpiar_candidato_campo(valor, "nombre")
    if not valor or len(valor) < 5 or len(valor) > 100:
        return False
    if re.search(r"\d|@|https?://", valor):
        return False
    tokens = re.findall(r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ'-]+", valor)
    if len(tokens) < 2 or len(tokens) > 8:
        return False
    norm = normalizar_etiqueta(valor)
    if norm in {normalizar_etiqueta(x) for x in _RUIDO_IDENTIDAD}:
        return False
    if any(
        ruido in norm
        for ruido in [
            "CERTIFICADO", "MEDICINA", "OCUPACIONAL", "EMPRESA", "INSTITUCION",
            "INSTITUCIÓN", "LABORATORIO", "CENTRO MEDICO", "CENTRO MÉDICO",
            "SALUD OCUPACIONAL", "FIRMA DEL", "NOMBRE DEL MEDICO", "NOMBRE DEL MÉDICO"
        ]
    ):
        return False
    letras = sum(c.isalpha() for c in valor)
    return letras / max(len(valor), 1) >= 0.65


def candidato_cargo_valido(valor):
    valor = limpiar_candidato_campo(valor, "cargo")
    if not valor or len(valor) < 3 or len(valor) > 120:
        return False
    norm = normalizar_etiqueta(valor)
    if norm in {normalizar_etiqueta(x) for x in _RUIDO_CARGO}:
        return False
    if any(
        ruido in norm
        for ruido in [
            "CERTIFICADO", "MEDICO OCUPACIONAL", "MÉDICO OCUPACIONAL",
            "FIRMA", "DOCUMENTO", "IDENTIFICACION", "IDENTIFICACIÓN",
            "FECHA DE NACIMIENTO", "GENERO", "GÉNERO"
        ]
    ):
        return False
    if re.fullmatch(r"[\d\s./-]+", valor):
        return False
    return sum(c.isalpha() for c in valor) >= 3


def candidato_lugar_valido(valor):
    valor = limpiar_candidato_campo(valor, "lugar")
    if not valor or len(valor) < 3 or len(valor) > 100:
        return False
    if re.search(r"https?://|www\.|@", valor, flags=re.IGNORECASE):
        return False
    if re.fullmatch(r"[\d\s:./-]+", valor):
        return False
    norm = normalizar_etiqueta(valor)
    if norm in {normalizar_etiqueta(x) for x in _RUIDO_LUGAR}:
        return False
    if any(
        ruido in norm
        for ruido in [
            "PAGINA", "CERTIFICADO", "LOGOTIPO", "FIRMA", "CONSENTIMIENTO",
            "HORA DE IMPRESION", "HORA DE IMPRESIÓN", "A.M", "P.M"
        ]
    ):
        return False
    return sum(c.isalpha() for c in valor) >= 3


def elegir_mejor_candidato(candidatos, tipo):
    validadores = {
        "nombre": candidato_nombre_valido,
        "cargo": candidato_cargo_valido,
        "lugar": candidato_lugar_valido,
    }
    validador = validadores[tipo]
    mejores = []

    for puntaje, valor, origen in candidatos:
        limpio = limpiar_candidato_campo(valor, tipo)
        if not validador(limpio):
            continue

        tokens = limpio.split()
        if tipo == "nombre":
            if 2 <= len(tokens) <= 5:
                puntaje += 8
            if limpio.upper() == limpio:
                puntaje += 3
        elif tipo == "cargo":
            if 1 <= len(tokens) <= 7:
                puntaje += 4
        elif tipo == "lugar":
            if 1 <= len(tokens) <= 5:
                puntaje += 4

        mejores.append((puntaje, -len(limpio), limpio, origen))

    if not mejores:
        return ""

    mejores.sort(reverse=True)
    valor = mejores[0][2]
    if tipo == "nombre":
        return valor.title()
    if tipo == "cargo":
        return corregir_ortografia_sst(valor).title()
    return valor.title()


def _coincide_etiqueta(columna, etiquetas):
    norm = normalizar_etiqueta(columna)
    return any(
        norm == normalizar_etiqueta(etiqueta)
        or normalizar_etiqueta(etiqueta) in norm
        for etiqueta in etiquetas
    )


def extraer_campo_por_etiquetas(lineas, etiquetas, tipo):
    """
    Obtiene candidatos en tres estructuras frecuentes:
    1. ETIQUETA: valor
    2. ETIQUETA    valor
    3. fila de encabezados + fila de valores.
    """
    candidatos = []
    etiquetas_ordenadas = sorted(etiquetas, key=len, reverse=True)
    patron_etiquetas = "|".join(re.escape(e) for e in etiquetas_ordenadas)

    for idx, linea in enumerate(lineas):
        if not linea or not linea.strip():
            continue

        # Valor en la misma línea después de dos puntos, guion o separación de tabla.
        m_inline = re.search(
            rf"(?:{patron_etiquetas})(?:\s*[:=]\s*|\s+-\s+|\s{{2,}}|\|)\s*(.+)$",
            linea,
            flags=re.IGNORECASE,
        )
        if m_inline:
            candidatos.append((115, m_inline.group(1), "etiqueta en línea"))

        columnas_header = dividir_columnas_estructuradas(linea)
        indices = [
            pos
            for pos, col in enumerate(columnas_header)
            if _coincide_etiqueta(col, etiquetas)
        ]

        if indices:
            for offset in range(1, 5):
                if idx + offset >= len(lineas):
                    break
                siguiente = lineas[idx + offset].strip()
                if not siguiente:
                    continue

                columnas_valor = dividir_columnas_estructuradas(siguiente)
                for pos in indices:
                    if pos < len(columnas_valor):
                        candidatos.append(
                            (105 - offset, columnas_valor[pos], "tabla encabezado/valor")
                        )

                # Si la etiqueta ocupa toda la línea, la línea siguiente completa
                # suele ser el dato, aunque el PDF haya perdido columnas.
                norm_linea = normalizar_etiqueta(linea)
                if any(
                    norm_linea == normalizar_etiqueta(etiqueta)
                    for etiqueta in etiquetas
                ):
                    candidatos.append((100 - offset, siguiente, "línea siguiente"))
                break

        # Etiqueta ubicada al inicio, pero sin separador reconocible.
        norm_linea = normalizar_etiqueta(linea)
        for etiqueta in etiquetas_ordenadas:
            norm_etiqueta = normalizar_etiqueta(etiqueta)
            if norm_linea.startswith(norm_etiqueta) and len(norm_linea) > len(norm_etiqueta):
                resto = linea[len(etiqueta):].strip(" |/-,_.:")
                if resto:
                    candidatos.append((92, resto, "etiqueta inicial"))
                break

    return elegir_mejor_candidato(candidatos, tipo)


def extraer_fecha_y_lugar_robusto(lineas, texto_completo):
    candidatos_lugar = []
    fechas = []

    patron_fecha_dmy = re.compile(
        r"\b(\d{1,2})\s*(?:[|/\-.]|\s)\s*(\d{1,2})\s*(?:[|/\-.]|\s)\s*(20\d{2})\b"
    )
    patron_fecha_ymd = re.compile(
        r"\b(20\d{2})\s*(?:[|/\-.]|\s)\s*(\d{1,2})\s*(?:[|/\-.]|\s)\s*(\d{1,2})\b"
    )

    # Tablas con encabezados DÍA | MES | AÑO | CIUDAD/MUNICIPIO.
    for idx, linea in enumerate(lineas):
        columnas = dividir_columnas_estructuradas(linea)
        columnas_norm = [normalizar_etiqueta(c) for c in columnas]

        ciudad_indices = [
            i for i, c in enumerate(columnas_norm)
            if any(k in c for k in ["CIUDAD", "MUNICIPIO", "LUGAR", "SEDE"])
        ]

        if ciudad_indices and any(
            k in " ".join(columnas_norm)
            for k in ["DIA", "MES", "ANO", "FECHA", "REALIZACION"]
        ):
            for offset in range(1, 4):
                if idx + offset >= len(lineas):
                    break
                valores = dividir_columnas_estructuradas(lineas[idx + offset])
                for pos in ciudad_indices:
                    if pos < len(valores):
                        candidatos_lugar.append(
                            (125 - offset, valores[pos], "tabla fecha/lugar")
                        )
                linea_valores = lineas[idx + offset]
                m_dmy = patron_fecha_dmy.search(linea_valores)
                m_ymd = patron_fecha_ymd.search(linea_valores)
                try:
                    if m_ymd:
                        fechas.append(
                            (
                                120 - offset,
                                datetime.date(
                                    int(m_ymd.group(1)),
                                    int(m_ymd.group(2)),
                                    int(m_ymd.group(3)),
                                ),
                            )
                        )
                    elif m_dmy:
                        fechas.append(
                            (
                                120 - offset,
                                datetime.date(
                                    int(m_dmy.group(3)),
                                    int(m_dmy.group(2)),
                                    int(m_dmy.group(1)),
                                ),
                            )
                        )
                except ValueError:
                    pass
                if valores:
                    break

        m_dmy = patron_fecha_dmy.search(linea)
        m_ymd = patron_fecha_ymd.search(linea)
        m_fecha = m_ymd or m_dmy
        if m_fecha:
            try:
                if m_ymd:
                    fecha_detectada = datetime.date(
                        int(m_ymd.group(1)),
                        int(m_ymd.group(2)),
                        int(m_ymd.group(3)),
                    )
                else:
                    fecha_detectada = datetime.date(
                        int(m_dmy.group(3)),
                        int(m_dmy.group(2)),
                        int(m_dmy.group(1)),
                    )
                fechas.append((90, fecha_detectada))
            except ValueError:
                pass

            antes = linea[:m_fecha.start()].strip(" |/-,_.:")
            despues = linea[m_fecha.end():].strip(" |/-,_.:")

            contexto = normalizar_etiqueta(linea)
            puntaje = 118 if any(
                k in contexto
                for k in ["REALIZACION", "CIUDAD", "MUNICIPIO", "LUGAR", "SEDE"]
            ) else 78

            if despues:
                candidatos_lugar.append((puntaje, despues, "después de fecha"))
            if antes and len(antes.split()) <= 7:
                # Evita usar el encabezado como ciudad.
                antes = re.sub(
                    r"(?i)\b(?:FECHA|CIUDAD|MUNICIPIO|LUGAR|REALIZACI[ÓO]N|DEL EXAMEN|DEL EXÁMEN)\b",
                    "",
                    antes,
                ).strip(" |/-,_.:")
                if antes:
                    candidatos_lugar.append((puntaje - 5, antes, "antes de fecha"))

    # Formatos como "Tunja, 15 de junio de 2026".
    meses = (
        "enero|febrero|marzo|abril|mayo|junio|julio|agosto|"
        "septiembre|octubre|noviembre|diciembre"
    )
    patron_ciudad_fecha = re.compile(
        rf"\b([A-Za-zÁÉÍÓÚÜÑáéíóúüñ][A-Za-zÁÉÍÓÚÜÑáéíóúüñ .'-]{{2,45}}),?\s+"
        rf"(\d{{1,2}}\s+de\s+(?:{meses})\s+de\s+20\d{{2}})\b",
        flags=re.IGNORECASE,
    )
    for m in patron_ciudad_fecha.finditer(texto_completo):
        lugar = m.group(1).strip()
        # Toma solo el segmento final si quedó unido a texto anterior.
        lugar = re.split(r"\n|[:;]", lugar)[-1].strip()
        candidatos_lugar.append((112, lugar, "ciudad y fecha en letras"))
        try:
            fechas.append((112, intentar_parsear_fecha(m.group(2))))
        except Exception:
            pass

    etiquetas_lugar = [
        "FECHA Y CIUDAD DE REALIZACIÓN",
        "FECHA Y CIUDAD DE REALIZACION",
        "CIUDAD DE REALIZACIÓN",
        "CIUDAD DE REALIZACION",
        "LUGAR DE REALIZACIÓN",
        "LUGAR DE REALIZACION",
        "MUNICIPIO DE REALIZACIÓN",
        "MUNICIPIO DE REALIZACION",
        "LUGAR DEL EXAMEN",
        "LUGAR DONDE SE REALIZÓ EL EXAMEN",
        "LUGAR DONDE SE REALIZO EL EXAMEN",
        "LUGAR DONDE SE REALIZARON LOS EXÁMENES",
        "LUGAR DONDE SE REALIZARON LOS EXAMENES",
        "CIUDAD DEL EXAMEN",
        "MUNICIPIO DEL EXAMEN",
        "SEDE DE ATENCIÓN",
        "SEDE DE ATENCION",
        "CENTRO MÉDICO",
        "CENTRO MEDICO",
        "IPS PRESTADORA",
        "IPS QUE REALIZA EL EXAMEN",
        "CIUDAD",
        "MUNICIPIO",
        "SEDE",
    ]
    lugar_etiquetado = extraer_campo_por_etiquetas(
        lineas, etiquetas_lugar, "lugar"
    )
    if lugar_etiquetado:
        candidatos_lugar.append((130, lugar_etiquetado, "etiqueta explícita"))

    lugar = elegir_mejor_candidato(candidatos_lugar, "lugar")
    fecha = max(fechas, key=lambda item: item[0])[1] if fechas else datetime.date.today()
    return fecha, lugar


def extraer_identidad_cargo_lugar(texto):
    lineas = [line.rstrip() for line in texto.splitlines()]
    etiquetas_nombre = [
        "APELLIDOS Y NOMBRES DEL TRABAJADOR",
        "NOMBRES Y APELLIDOS DEL TRABAJADOR",
        "APELLIDOS Y NOMBRES",
        "NOMBRES Y APELLIDOS",
        "NOMBRE DEL TRABAJADOR",
        "NOMBRE TRABAJADOR",
        "NOMBRE COMPLETO",
        "NOMBRES COMPLETOS",
        "NOMBRE(S) Y APELLIDO(S)",
        "APELLIDO(S) Y NOMBRE(S)",
        "NOMBRE Y APELLIDOS",
        "NOMBRES DEL TRABAJADOR",
        "TRABAJADOR",
        "PACIENTE",
    ]
    etiquetas_cargo = [
        "CARGO ACTUAL DEL TRABAJADOR",
        "CARGO DEL TRABAJADOR",
        "CARGO ACTUAL",
        "OCUPACIÓN ACTUAL",
        "OCUPACION ACTUAL",
        "OCUPACIÓN",
        "OCUPACION",
        "PUESTO DE TRABAJO",
        "OCUPACIÓN DEL TRABAJADOR",
        "OCUPACION DEL TRABAJADOR",
        "CARGO U OCUPACIÓN",
        "CARGO U OCUPACION",
        "CARGO / OCUPACIÓN",
        "CARGO / OCUPACION",
        "PUESTO",
        "OFICIO",
        "LABOR",
        "CARGO",
    ]

    nombre = extraer_campo_por_etiquetas(
        lineas, etiquetas_nombre, "nombre"
    )
    cargo = extraer_campo_por_etiquetas(
        lineas, etiquetas_cargo, "cargo"
    )
    fecha, lugar = extraer_fecha_y_lugar_robusto(
        lineas, texto
    )

    return {
        "nombre": nombre,
        "cargo": cargo,
        "fecha": fecha,
        "lugar": lugar,
    }



# --- EXTRACTOR ESPECIALIZADO PARA LOS DOS FORMATOS DE CERTIFICADO ---
# Se añade sobre la base existente. No reemplaza el analizador de exámenes,
# recomendaciones, PVE, observaciones, remisiones ni generación de archivos.

_ETIQUETAS_NOMBRE_CERTIFICADO = [
    "NOMBRES Y APELLIDOS TRABAJADOR",
    "NOMBRES Y APELLIDOS DEL TRABAJADOR",
    "APELLIDOS Y NOMBRES TRABAJADOR",
    "APELLIDOS Y NOMBRES DEL TRABAJADOR",
    "NOMBRES Y APELLIDOS",
    "APELLIDOS Y NOMBRES",
    "NOMBRE DEL TRABAJADOR",
    "NOMBRE COMPLETO",
    "PACIENTE",
]

_ETIQUETAS_CARGO_CERTIFICADO = [
    "CARGO DEL TRABAJADOR",
    "CARGO ACTUAL DEL TRABAJADOR",
    "CARGO ACTUAL",
    "CARGO U OCUPACIÓN",
    "CARGO U OCUPACION",
    "OCUPACIÓN DEL TRABAJADOR",
    "OCUPACION DEL TRABAJADOR",
    "PUESTO DE TRABAJO",
    "OCUPACIÓN",
    "OCUPACION",
    "OFICIO",
    "LABOR",
    "CARGO",
]

_ETIQUETAS_FECHA_CERTIFICADO = [
    "FECHA DE REALIZACIÓN DEL EXAMEN",
    "FECHA DE REALIZACION DEL EXAMEN",
    "FECHA DE REALIZACIÓN DE LOS EXÁMENES",
    "FECHA DE REALIZACION DE LOS EXAMENES",
    "FECHA DEL EXAMEN",
    "FECHA EXAMEN",
    "FECHA DE ATENCIÓN",
    "FECHA DE ATENCION",
]

_ETIQUETAS_LUGAR_CERTIFICADO = [
    "FECHA Y CIUDAD DE REALIZACIÓN",
    "FECHA Y CIUDAD DE REALIZACION",
    "CIUDAD DE REALIZACIÓN DEL EXAMEN",
    "CIUDAD DE REALIZACION DEL EXAMEN",
    "LUGAR DE REALIZACIÓN DEL EXAMEN",
    "LUGAR DE REALIZACION DEL EXAMEN",
    "LUGAR DE REALIZACIÓN DE LOS EXÁMENES",
    "LUGAR DE REALIZACION DE LOS EXAMENES",
    "LUGAR DONDE SE REALIZARON LOS EXÁMENES",
    "LUGAR DONDE SE REALIZARON LOS EXAMENES",
    "MUNICIPIO DE REALIZACIÓN",
    "MUNICIPIO DE REALIZACION",
    "CIUDAD DEL EXAMEN",
    "MUNICIPIO DEL EXAMEN",
    "LUGAR DEL EXAMEN",
    "SEDE DE ATENCIÓN",
    "SEDE DE ATENCION",
    "CIUDAD",
    "MUNICIPIO",
    "LUGAR",
    "SEDE",
]

_ETIQUETAS_IPS_CERTIFICADO = [
    "IPS QUE REALIZA EL EXAMEN",
    "IPS PRESTADORA",
    "CENTRO MÉDICO",
    "CENTRO MEDICO",
    "INSTITUCIÓN PRESTADORA",
    "INSTITUCION PRESTADORA",
]

_PATRONES_LEGALES_RECOMENDACIONES = [
    r"\bconsentimiento(?:\s+informado)?\b",
    r"\bautorizo\b",
    r"\bautorización\s+para\s+el\s+tratamiento\s+de\s+datos\b",
    r"\bautorizacion\s+para\s+el\s+tratamiento\s+de\s+datos\b",
    r"\btratamiento\s+de\s+datos(?:\s+personales)?\b",
    r"\bprotección\s+de\s+datos\b",
    r"\bproteccion\s+de\s+datos\b",
    r"\bhabeas\s+data\b",
    r"\bley\s+1581\b",
    r"\bdeclaro\b",
    r"\bmanifiesto\b",
    r"\bhe\s+sido\s+informad[oa]\b",
    r"\bacepto\s+(?:el|la|los|las)\b",
    r"\bconstancia\b",
    r"\briesgos\s+y\s+beneficios\b",
    r"\bfirma\s+(?:del|de\s+la)\s+(?:trabajador|paciente|usuario|evaluado)\b",
    r"\bfirma\s+del\s+m[eé]dico\b",
    r"\bhuella\b",
    r"\bdocumento\s+de\s+identidad\b",
    r"\bresponsabilidad\s+del\s+paciente\b",
    r"\bdeclaración\s+del\s+paciente\b",
    r"\bdeclaracion\s+del\s+paciente\b",
    r"\bderechos\s+y\s+deberes\b",
    r"\binformación\s+suministrada\s+es\s+verdadera\b",
    r"\binformacion\s+suministrada\s+es\s+verdadera\b",
]

_ENCABEZADOS_LEGALES = [
    "CONSENTIMIENTO INFORMADO",
    "CONSENTIMIENTO",
    "AUTORIZACIÓN PARA TRATAMIENTO DE DATOS",
    "AUTORIZACION PARA TRATAMIENTO DE DATOS",
    "TRATAMIENTO DE DATOS PERSONALES",
    "DECLARACIÓN DEL PACIENTE",
    "DECLARACION DEL PACIENTE",
    "AUTORIZO",
    "CONSTANCIA",
    "FIRMA DEL TRABAJADOR",
    "FIRMA DEL PACIENTE",
    "FIRMA DEL USUARIO",
    "HUELLA",
    "HABEAS DATA",
]


def es_contenido_legal_recomendacion(texto):
    if not texto:
        return False
    limpio = re.sub(r"\s+", " ", str(texto)).strip()
    return any(
        re.search(patron, limpio, flags=re.IGNORECASE)
        for patron in _PATRONES_LEGALES_RECOMENDACIONES
    )


def es_encabezado_legal(texto):
    if not texto:
        return False
    normalizado = normalizar_etiqueta(texto)
    return any(
        encabezado in normalizado
        for encabezado in _ENCABEZADOS_LEGALES
    )


def recortar_contenido_legal(texto):
    """
    Corta una recomendación justo antes del primer consentimiento,
    autorización, firma o declaración legal.
    """
    if not texto:
        return ""

    texto = str(texto)
    posiciones = []
    for patron in _PATRONES_LEGALES_RECOMENDACIONES:
        coincidencia = re.search(patron, texto, flags=re.IGNORECASE)
        if coincidencia:
            posiciones.append(coincidencia.start())

    if posiciones:
        texto = texto[:min(posiciones)]

    lineas_validas = []
    for linea in texto.splitlines():
        if es_encabezado_legal(linea) or es_contenido_legal_recomendacion(linea):
            break
        lineas_validas.append(linea)

    return re.sub(r"\s+", " ", " ".join(lineas_validas)).strip(" .;:-_/|")


def filtrar_recomendaciones_clinicas(recomendaciones):
    """
    Filtro final obligatorio. No permite que consentimientos, autorizaciones,
    firmas, habeas data o declaraciones entren al Word.
    """
    resultado = []
    vistos = set()

    for recomendacion in recomendaciones or []:
        limpia = recortar_contenido_legal(recomendacion)

        limpia = re.sub(
            r"^(?:Audiometría|Espirometría|Optometría|Visiometría|"
            r"Examen Clínico Ocupacional|Énfasis Osteomuscular|"
            r"Electrocardiograma|Frotis|Cuadro Hemático|Colesterol|"
            r"Triglicéridos|Parcial de Orina|VSH|PCR)\s*:\s*$",
            "",
            limpia,
            flags=re.IGNORECASE,
        ).strip()

        if (
            not limpia
            or es_vacio_o_estado(limpia)
            or es_contenido_legal_recomendacion(limpia)
        ):
            continue

        clave = normalizar_etiqueta(limpia)
        if clave not in vistos:
            vistos.add(clave)
            resultado.append(limpia)

    return resultado


def _limpiar_celda_certificado(valor):
    if valor is None:
        return ""
    valor = str(valor).replace("\x00", " ")
    valor = re.sub(r"[ \t]+", " ", valor)
    valor = re.sub(r"\s*\n\s*", "\n", valor)
    return valor.strip(" |/-,_.:")


def _normalizar_lista_celdas(fila):
    return [
        _limpiar_celda_certificado(celda)
        for celda in (fila or [])
    ]


def _contiene_alguna_etiqueta(texto, etiquetas):
    normalizado = normalizar_etiqueta(texto)
    return any(
        normalizar_etiqueta(etiqueta) in normalizado
        for etiqueta in etiquetas
    )


def _es_rotulo_general(texto):
    normalizado = normalizar_etiqueta(texto)
    rotulos = (
        _ETIQUETAS_NOMBRE_CERTIFICADO
        + _ETIQUETAS_CARGO_CERTIFICADO
        + _ETIQUETAS_FECHA_CERTIFICADO
        + _ETIQUETAS_LUGAR_CERTIFICADO
        + _ETIQUETAS_IPS_CERTIFICADO
        + [
            "DOCUMENTO",
            "IDENTIFICACIÓN",
            "IDENTIFICACION",
            "CÉDULA",
            "CEDULA",
            "EDAD",
            "GÉNERO",
            "GENERO",
            "SEXO",
            "EMPRESA",
            "EPS",
            "ARL",
            "AFP",
            "DÍA",
            "DIA",
            "MES",
            "AÑO",
            "ANO",
        ]
    )
    return any(
        normalizado == normalizar_etiqueta(rotulo)
        or normalizado.startswith(normalizar_etiqueta(rotulo) + " ")
        for rotulo in rotulos
    )


def _extraer_valor_inline(celda, etiquetas):
    celda = _limpiar_celda_certificado(celda)
    if not celda:
        return ""

    # Admite "CARGO: CONDUCTOR" y celdas con salto de línea.
    plano = re.sub(r"\s*\n\s*", " | ", celda)
    for etiqueta in sorted(etiquetas, key=len, reverse=True):
        patron = re.compile(
            rf"^\s*{re.escape(etiqueta)}\s*(?:[:=\-|]\s*)?(.+)$",
            flags=re.IGNORECASE,
        )
        coincidencia = patron.match(plano)
        if coincidencia:
            resto = coincidencia.group(1).strip(" |/-,_.:")
            if resto and not _es_rotulo_general(resto):
                return resto

        lineas = [
            linea.strip(" |/-,_.:")
            for linea in celda.splitlines()
            if linea.strip(" |/-,_.:")
        ]
        if (
            lineas
            and normalizar_etiqueta(lineas[0])
            == normalizar_etiqueta(etiqueta)
            and len(lineas) > 1
        ):
            resto = " ".join(lineas[1:]).strip(" |/-,_.:")
            if resto and not _es_rotulo_general(resto):
                return resto

    return ""


def _nombre_muy_valido(valor):
    limpio = limpiar_candidato_campo(valor, "nombre")
    if not candidato_nombre_valido(limpio):
        return False

    norm = normalizar_etiqueta(limpio)
    prohibidas = [
        "FECHA",
        "REALIZACION",
        "EXAMEN",
        "CARGO",
        "EMPRESA",
        "IDENTIFICACION",
        "DOCUMENTO",
        "CERTIFICADO",
        "CONCEPTO",
        "RECOMENDACION",
        "VIGILANCIA",
        "CONSENTIMIENTO",
        "TRABAJADOR",
        "APELLIDOS Y NOMBRES",
        "NOMBRES Y APELLIDOS",
    ]
    if any(palabra in norm for palabra in prohibidas):
        return False

    tokens = re.findall(
        r"[A-Za-zÁÉÍÓÚÜÑáéíóúüñ'-]+",
        limpio,
    )
    return 2 <= len(tokens) <= 7


def _cargo_muy_valido(valor):
    limpio = limpiar_candidato_campo(valor, "cargo")
    if not candidato_cargo_valido(limpio):
        return False

    norm = normalizar_etiqueta(limpio)
    prohibidas = [
        "FECHA",
        "REALIZACION",
        "EXAMEN",
        "DOCUMENTO",
        "IDENTIFICACION",
        "NOMBRES Y APELLIDOS",
        "APELLIDOS Y NOMBRES",
        "CERTIFICADO",
        "CONSENTIMIENTO",
        "DIA MES ANO",
    ]
    return not any(palabra in norm for palabra in prohibidas)


def _lugar_muy_valido(valor):
    limpio = limpiar_candidato_campo(valor, "lugar")
    if not candidato_lugar_valido(limpio):
        return False

    norm = normalizar_etiqueta(limpio)
    prohibidas = [
        "FECHA",
        "REALIZACION",
        "EXAMEN",
        "DIA",
        "MES",
        "ANO",
        "DOCUMENTO",
        "IDENTIFICACION",
        "CONSENTIMIENTO",
    ]
    return not any(palabra == norm for palabra in prohibidas)


def _agregar_candidato_especial(
    candidatos,
    tipo,
    puntaje,
    valor,
    origen,
):
    if not valor:
        return

    validadores = {
        "nombre": _nombre_muy_valido,
        "cargo": _cargo_muy_valido,
        "lugar": _lugar_muy_valido,
    }
    limpio = limpiar_candidato_campo(valor, tipo)
    if validadores[tipo](limpio):
        candidatos[tipo].append(
            (puntaje, limpio, origen)
        )


def _buscar_valor_debajo(
    filas,
    fila_inicio,
    columna,
    tipo,
    max_filas=5,
):
    """
    Formato 1 de las imágenes:
    ETIQUETA
    VALOR

    También admite celdas combinadas, porque revisa la misma columna,
    celdas cercanas y la fila completa.
    """
    validadores = {
        "nombre": _nombre_muy_valido,
        "cargo": _cargo_muy_valido,
        "lugar": _lugar_muy_valido,
    }
    validador = validadores[tipo]

    for salto in range(1, max_filas + 1):
        indice = fila_inicio + salto
        if indice >= len(filas):
            break

        fila = filas[indice]
        if not any(fila):
            continue

        opciones = []

        if columna < len(fila):
            opciones.append(fila[columna])

        # Las tablas con celdas fusionadas pueden desplazar el valor una columna.
        for desplazamiento in (-2, -1, 1, 2):
            pos = columna + desplazamiento
            if 0 <= pos < len(fila):
                opciones.append(fila[pos])

        # Si la fila contiene un único texto claro, se considera el valor
        # correspondiente al encabezado superior.
        no_vacias = [
            celda for celda in fila
            if celda and not _es_rotulo_general(celda)
        ]
        if len(no_vacias) == 1:
            opciones.extend(no_vacias)

        for opcion in opciones:
            limpio = limpiar_candidato_campo(opcion, tipo)
            if validador(limpio):
                return limpio, salto

    return "", 0


def _buscar_valor_derecha(fila, columna, tipo):
    """
    Formato 2 de las imágenes:
    ETIQUETA | VALOR
    """
    validadores = {
        "nombre": _nombre_muy_valido,
        "cargo": _cargo_muy_valido,
        "lugar": _lugar_muy_valido,
    }
    validador = validadores[tipo]

    for pos in range(columna + 1, min(len(fila), columna + 5)):
        celda = fila[pos]
        if not celda or _es_rotulo_general(celda):
            continue
        limpio = limpiar_candidato_campo(celda, tipo)
        if validador(limpio):
            return limpio, pos - columna

    return "", 0


def _fecha_desde_componentes(dia, mes, anio):
    try:
        dia = int(str(dia).strip())
        mes = int(str(mes).strip())
        anio = int(str(anio).strip())
        if anio < 100:
            anio += 2000
        return datetime.date(anio, mes, dia)
    except (TypeError, ValueError):
        return None


def _buscar_fecha_en_texto(texto):
    if not texto:
        return None

    texto = re.sub(r"\s+", " ", str(texto))

    patrones = [
        (
            re.compile(
                r"\b(20\d{2})\s*[-/.\s]\s*(\d{1,2})\s*[-/.\s]\s*(\d{1,2})\b"
            ),
            lambda m: _fecha_desde_componentes(
                m.group(3),
                m.group(2),
                m.group(1),
            ),
        ),
        (
            re.compile(
                r"\b(\d{1,2})\s*[-/.\s]\s*(\d{1,2})\s*[-/.\s]\s*(20\d{2})\b"
            ),
            lambda m: _fecha_desde_componentes(
                m.group(1),
                m.group(2),
                m.group(3),
            ),
        ),
    ]

    for patron, constructor in patrones:
        coincidencia = patron.search(texto)
        if coincidencia:
            fecha = constructor(coincidencia)
            if fecha:
                return fecha

    meses = {
        "enero": 1,
        "febrero": 2,
        "marzo": 3,
        "abril": 4,
        "mayo": 5,
        "junio": 6,
        "julio": 7,
        "agosto": 8,
        "septiembre": 9,
        "octubre": 10,
        "noviembre": 11,
        "diciembre": 12,
    }
    coincidencia = re.search(
        r"\b(\d{1,2})\s+de\s+([a-záéíóúüñ]+)\s+de\s+(20\d{2})\b",
        texto,
        flags=re.IGNORECASE,
    )
    if coincidencia:
        mes = meses.get(coincidencia.group(2).lower())
        if mes:
            return _fecha_desde_componentes(
                coincidencia.group(1),
                mes,
                coincidencia.group(3),
            )

    return None


def _buscar_fecha_tabla(filas, fila_inicio, columna_inicio=0):
    """
    Reconoce ambos diseños:
    - FECHA DE REALIZACIÓN DEL EXAMEN | DÍA | MES | AÑO / 04 | 06 | 2026
    - FECHA DE REALIZACIÓN DEL EXAMEN | 04/06/2026
    """
    limite = min(len(filas), fila_inicio + 7)

    # Primero intenta una fecha completa en las filas cercanas.
    for indice in range(fila_inicio, limite):
        fila = filas[indice]
        texto_fila = " | ".join(celda for celda in fila if celda)
        fecha = _buscar_fecha_en_texto(texto_fila)
        if fecha:
            return fecha, 250 - (indice - fila_inicio)

    # Luego localiza subencabezados DÍA, MES y AÑO.
    for indice in range(fila_inicio, limite):
        fila = filas[indice]
        normalizadas = [
            normalizar_etiqueta(celda)
            for celda in fila
        ]

        indice_dia = next(
            (
                pos for pos, valor in enumerate(normalizadas)
                if valor in {"DIA", "DÍA"}
            ),
            None,
        )
        indice_mes = next(
            (
                pos for pos, valor in enumerate(normalizadas)
                if valor == "MES"
            ),
            None,
        )
        indice_anio = next(
            (
                pos for pos, valor in enumerate(normalizadas)
                if valor in {"ANO", "AÑO"}
            ),
            None,
        )

        if None not in (
            indice_dia,
            indice_mes,
            indice_anio,
        ):
            for salto in range(1, 4):
                fila_valor_idx = indice + salto
                if fila_valor_idx >= len(filas):
                    break

                valores = filas[fila_valor_idx]
                if max(
                    indice_dia,
                    indice_mes,
                    indice_anio,
                ) >= len(valores):
                    continue

                fecha = _fecha_desde_componentes(
                    valores[indice_dia],
                    valores[indice_mes],
                    valores[indice_anio],
                )
                if fecha:
                    return fecha, 270 - salto

    # Respaldo: tres números consecutivos cerca del encabezado.
    for indice in range(fila_inicio + 1, limite):
        numeros = []
        for celda in filas[indice]:
            if re.fullmatch(r"\d{1,4}", celda or ""):
                numeros.append(int(celda))

        for pos in range(0, len(numeros) - 2):
            trio = numeros[pos:pos + 3]

            if 1 <= trio[0] <= 31 and 1 <= trio[1] <= 12 and 2000 <= trio[2] <= 2100:
                fecha = _fecha_desde_componentes(
                    trio[0],
                    trio[1],
                    trio[2],
                )
                if fecha:
                    return fecha, 230

            if 2000 <= trio[0] <= 2100 and 1 <= trio[1] <= 12 and 1 <= trio[2] <= 31:
                fecha = _fecha_desde_componentes(
                    trio[2],
                    trio[1],
                    trio[0],
                )
                if fecha:
                    return fecha, 230

    return None, 0


def _quitar_fecha_para_lugar(texto):
    if not texto:
        return ""

    valor = str(texto)
    valor = re.sub(
        r"\b20\d{2}\s*[-/.\s]\s*\d{1,2}\s*[-/.\s]\s*\d{1,2}\b",
        " ",
        valor,
    )
    valor = re.sub(
        r"\b\d{1,2}\s*[-/.\s]\s*\d{1,2}\s*[-/.\s]\s*20\d{2}\b",
        " ",
        valor,
    )
    valor = re.sub(
        r"\b\d{1,2}\s+de\s+[a-záéíóúüñ]+\s+de\s+20\d{2}\b",
        " ",
        valor,
        flags=re.IGNORECASE,
    )
    valor = re.sub(
        r"\b(?:FECHA|DÍA|DIA|MES|AÑO|ANO|CIUDAD|MUNICIPIO|LUGAR|"
        r"REALIZACIÓN|REALIZACION|DEL EXAMEN|DE LOS EXÁMENES|"
        r"DE LOS EXAMENES|SEDE)\b",
        " ",
        valor,
        flags=re.IGNORECASE,
    )
    return re.sub(r"\s+", " ", valor).strip(" |/-,_.:")


def _extraer_lineas_por_coordenadas(page):
    """
    Reconstruye líneas respetando la posición horizontal de las palabras.
    Conserva separaciones amplias para que el analizador detecte columnas.
    """
    try:
        palabras = page.extract_words(
            x_tolerance=2,
            y_tolerance=3,
            keep_blank_chars=False,
            use_text_flow=False,
        ) or []
    except Exception:
        return []

    grupos = []
    for palabra in sorted(
        palabras,
        key=lambda item: (
            round(float(item["top"]), 1),
            float(item["x0"]),
        ),
    ):
        top = float(palabra["top"])
        grupo = None

        for existente in reversed(grupos[-10:]):
            if abs(existente["top"] - top) <= 3.2:
                grupo = existente
                break

        if grupo is None:
            grupo = {
                "top": top,
                "words": [],
            }
            grupos.append(grupo)

        grupo["words"].append(palabra)

    lineas = []
    for grupo in grupos:
        palabras_linea = sorted(
            grupo["words"],
            key=lambda item: float(item["x0"]),
        )
        partes = []
        x1_anterior = None

        for palabra in palabras_linea:
            if x1_anterior is not None:
                separacion = float(palabra["x0"]) - x1_anterior
                partes.append(
                    "    " if separacion > 16 else " "
                )
            partes.append(str(palabra["text"]))
            x1_anterior = float(palabra["x1"])

        linea = "".join(partes).strip()
        if linea:
            lineas.append(linea)

    return lineas


def _ocr_pagina_si_disponible(page):
    """
    OCR complementario para certificados escaneados.
    Si Tesseract no está instalado, la aplicación continúa con pdfplumber.
    """
    if not OCR_DISPONIBLE:
        return ""

    try:
        imagen = page.to_image(
            resolution=230
        ).original.convert("L")
        imagen = ImageOps.autocontrast(imagen)
        imagen = ImageEnhance.Contrast(
            imagen
        ).enhance(1.35)
        imagen = imagen.filter(
            ImageFilter.SHARPEN
        )

        texto = pytesseract.image_to_string(
            imagen,
            lang="spa+eng",
            config="--oem 1 --psm 6",
            timeout=80,
        )
        return re.sub(
            r"\n{3,}",
            "\n\n",
            texto or "",
        ).strip()
    except Exception:
        return ""


def _extraer_de_filas_certificado(
    filas,
    candidatos,
    fechas,
    bonus=0,
    origen="tabla",
):
    """
    Procesa los dos formatos de las imágenes sin alterar los demás datos.
    """
    for fila_idx, fila in enumerate(filas):
        if not any(fila):
            continue

        for columna, celda in enumerate(fila):
            if not celda:
                continue

            # NOMBRE: valor en la misma celda, a la derecha o debajo.
            if _contiene_alguna_etiqueta(
                celda,
                _ETIQUETAS_NOMBRE_CERTIFICADO,
            ):
                inline = _extraer_valor_inline(
                    celda,
                    _ETIQUETAS_NOMBRE_CERTIFICADO,
                )
                _agregar_candidato_especial(
                    candidatos,
                    "nombre",
                    430 + bonus,
                    inline,
                    f"{origen}: nombre en misma celda",
                )

                derecha, distancia = _buscar_valor_derecha(
                    fila,
                    columna,
                    "nombre",
                )
                _agregar_candidato_especial(
                    candidatos,
                    "nombre",
                    420 - distancia + bonus,
                    derecha,
                    f"{origen}: nombre a la derecha",
                )

                debajo, salto = _buscar_valor_debajo(
                    filas,
                    fila_idx,
                    columna,
                    "nombre",
                )
                _agregar_candidato_especial(
                    candidatos,
                    "nombre",
                    425 - salto + bonus,
                    debajo,
                    f"{origen}: nombre debajo",
                )

            # CARGO: mismo esquema.
            if _contiene_alguna_etiqueta(
                celda,
                _ETIQUETAS_CARGO_CERTIFICADO,
            ):
                inline = _extraer_valor_inline(
                    celda,
                    _ETIQUETAS_CARGO_CERTIFICADO,
                )
                _agregar_candidato_especial(
                    candidatos,
                    "cargo",
                    420 + bonus,
                    inline,
                    f"{origen}: cargo en misma celda",
                )

                derecha, distancia = _buscar_valor_derecha(
                    fila,
                    columna,
                    "cargo",
                )
                _agregar_candidato_especial(
                    candidatos,
                    "cargo",
                    410 - distancia + bonus,
                    derecha,
                    f"{origen}: cargo a la derecha",
                )

                debajo, salto = _buscar_valor_debajo(
                    filas,
                    fila_idx,
                    columna,
                    "cargo",
                )
                _agregar_candidato_especial(
                    candidatos,
                    "cargo",
                    415 - salto + bonus,
                    debajo,
                    f"{origen}: cargo debajo",
                )

            # FECHA: detecta fecha completa o DÍA/MES/AÑO en filas siguientes.
            if _contiene_alguna_etiqueta(
                celda,
                _ETIQUETAS_FECHA_CERTIFICADO,
            ):
                fecha_inline = _buscar_fecha_en_texto(celda)
                if fecha_inline:
                    fechas.append(
                        (
                            440 + bonus,
                            fecha_inline,
                            f"{origen}: fecha misma celda",
                        )
                    )

                fecha_tabla, puntaje = _buscar_fecha_tabla(
                    filas,
                    fila_idx,
                    columna,
                )
                if fecha_tabla:
                    fechas.append(
                        (
                            puntaje + bonus,
                            fecha_tabla,
                            f"{origen}: fecha por componentes",
                        )
                    )

            # LUGAR: valor en línea, derecha o debajo.
            if _contiene_alguna_etiqueta(
                celda,
                _ETIQUETAS_LUGAR_CERTIFICADO,
            ):
                inline = _quitar_fecha_para_lugar(
                    _extraer_valor_inline(
                        celda,
                        _ETIQUETAS_LUGAR_CERTIFICADO,
                    )
                )
                _agregar_candidato_especial(
                    candidatos,
                    "lugar",
                    435 + bonus,
                    inline,
                    f"{origen}: lugar misma celda",
                )

                derecha, distancia = _buscar_valor_derecha(
                    fila,
                    columna,
                    "lugar",
                )
                _agregar_candidato_especial(
                    candidatos,
                    "lugar",
                    425 - distancia + bonus,
                    _quitar_fecha_para_lugar(derecha),
                    f"{origen}: lugar a la derecha",
                )

                debajo, salto = _buscar_valor_debajo(
                    filas,
                    fila_idx,
                    columna,
                    "lugar",
                )
                _agregar_candidato_especial(
                    candidatos,
                    "lugar",
                    430 - salto + bonus,
                    _quitar_fecha_para_lugar(debajo),
                    f"{origen}: lugar debajo",
                )

            # La IPS es respaldo, no desplaza una ciudad explícita.
            if _contiene_alguna_etiqueta(
                celda,
                _ETIQUETAS_IPS_CERTIFICADO,
            ):
                inline = _extraer_valor_inline(
                    celda,
                    _ETIQUETAS_IPS_CERTIFICADO,
                )
                _agregar_candidato_especial(
                    candidatos,
                    "lugar",
                    260 + bonus,
                    inline,
                    f"{origen}: IPS explícita",
                )


def _filas_desde_lineas(lineas):
    """
    Convierte texto visual/OCR en filas para reutilizar el mismo extractor.
    Las separaciones amplias y barras se interpretan como columnas.
    """
    filas = []
    for linea in lineas:
        linea = str(linea).strip()
        if not linea:
            continue

        columnas = [
            _limpiar_celda_certificado(columna)
            for columna in re.split(
                r"\s{3,}|\t+|\|",
                linea,
            )
        ]
        columnas = [
            columna
            for columna in columnas
            if columna
        ]
        if columnas:
            filas.append(columnas)

    return filas


def extraer_metadatos_pdf_estructurados(
    pdf_raw_data,
    texto_completo="",
):
    """
    Extractor prioritario de nombre, cargo, fecha y lugar.

    Orden de confianza:
    1. tablas reales;
    2. coordenadas de palabras;
    3. OCR;
    4. texto consolidado existente.

    No modifica ninguna otra función o campo del Portal SST.
    """
    candidatos = {
        "nombre": [],
        "cargo": [],
        "lugar": [],
    }
    fechas = []

    try:
        documento = pdfplumber.open(
            io.BytesIO(pdf_raw_data)
        )
    except Exception:
        documento = None

    if documento is not None:
        with documento as p_file:
            for numero_pagina, page in enumerate(
                p_file.pages,
                start=1,
            ):
                bonus = max(
                    0,
                    35 - (numero_pagina - 1) * 8,
                )

                # A. Tablas detectadas por bordes/celdas.
                try:
                    tablas = page.extract_tables() or []
                except Exception:
                    tablas = []

                for tabla in tablas:
                    filas = [
                        _normalizar_lista_celdas(fila)
                        for fila in (tabla or [])
                    ]
                    _extraer_de_filas_certificado(
                        filas,
                        candidatos,
                        fechas,
                        bonus=bonus,
                        origen=(
                            f"tabla página {numero_pagina}"
                        ),
                    )

                # B. Formato visual sin bordes detectables.
                lineas_coordenadas = (
                    _extraer_lineas_por_coordenadas(page)
                )
                filas_coordenadas = _filas_desde_lineas(
                    lineas_coordenadas
                )
                _extraer_de_filas_certificado(
                    filas_coordenadas,
                    candidatos,
                    fechas,
                    bonus=bonus - 20,
                    origen=(
                        f"coordenadas página {numero_pagina}"
                    ),
                )

                # C. OCR, especialmente para imágenes escaneadas.
                texto_base = page.extract_text(
                    layout=False
                ) or ""
                requiere_ocr = (
                    numero_pagina <= 2
                    or len(
                        re.sub(
                            r"\s+",
                            "",
                            texto_base,
                        )
                    ) < 150
                )
                if requiere_ocr:
                    texto_ocr = _ocr_pagina_si_disponible(
                        page
                    )
                    if texto_ocr:
                        filas_ocr = _filas_desde_lineas(
                            texto_ocr.splitlines()
                        )
                        _extraer_de_filas_certificado(
                            filas_ocr,
                            candidatos,
                            fechas,
                            bonus=bonus - 35,
                            origen=(
                                f"OCR página {numero_pagina}"
                            ),
                        )

                        fecha_ocr = _buscar_fecha_en_texto(
                            texto_ocr
                        )
                        if fecha_ocr:
                            fechas.append(
                                (
                                    250 + bonus,
                                    fecha_ocr,
                                    (
                                        f"OCR página "
                                        f"{numero_pagina}"
                                    ),
                                )
                            )

    # D. Analizador textual preexistente como respaldo.
    if texto_completo:
        respaldo = extraer_identidad_cargo_lugar(
            texto_completo
        )

        _agregar_candidato_especial(
            candidatos,
            "nombre",
            170,
            respaldo.get("nombre", ""),
            "analizador textual existente",
        )
        _agregar_candidato_especial(
            candidatos,
            "cargo",
            165,
            respaldo.get("cargo", ""),
            "analizador textual existente",
        )
        _agregar_candidato_especial(
            candidatos,
            "lugar",
            175,
            respaldo.get("lugar", ""),
            "analizador textual existente",
        )

        if respaldo.get("fecha"):
            fechas.append(
                (
                    165,
                    respaldo["fecha"],
                    "analizador textual existente",
                )
            )

    resultado = {
        "nombre": elegir_mejor_candidato(
            candidatos["nombre"],
            "nombre",
        ),
        "cargo": elegir_mejor_candidato(
            candidatos["cargo"],
            "cargo",
        ),
        "lugar": elegir_mejor_candidato(
            candidatos["lugar"],
            "lugar",
        ),
    }

    fechas_validas = [
        (puntaje, fecha, origen)
        for puntaje, fecha, origen in fechas
        if isinstance(fecha, datetime.date)
        and 2000 <= fecha.year <= 2100
    ]
    if fechas_validas:
        resultado["fecha"] = max(
            fechas_validas,
            key=lambda item: item[0],
        )[1]

    return resultado




def extraer_texto_pdf_robusto(pdf_raw_data):
    """
    Conserva la lectura original y añade:
    - distribución visual;
    - líneas reconstruidas con coordenadas;
    - filas de tablas;
    - OCR complementario en las dos primeras páginas y cuando no hay texto.

    No altera los demás campos ni el flujo de generación.
    """
    lineas_salida = []
    vistos = set()

    def agregar(fragmento, permitir_repetido=False):
        if not fragmento:
            return
        for linea in str(fragmento).splitlines():
            linea = linea.replace("\t", "    ").strip()
            if not linea:
                continue
            clave = normalizar_etiqueta(linea)
            if not clave:
                continue
            if permitir_repetido or clave not in vistos:
                vistos.add(clave)
                lineas_salida.append(linea)

    with pdfplumber.open(io.BytesIO(pdf_raw_data)) as p_file:
        for numero_pagina, page in enumerate(p_file.pages, start=1):
            texto_layout = page.extract_text(
                x_tolerance=2,
                y_tolerance=3,
                layout=True,
            ) or ""
            texto_normal = page.extract_text(
                x_tolerance=2,
                y_tolerance=3,
                layout=False,
            ) or ""

            agregar(texto_layout)
            agregar(texto_normal)

            # Reconstrucción por coordenadas: especialmente útil para
            # encabezados NOMBRE | CARGO y su fila de valores.
            for linea in _extraer_lineas_por_coordenadas(page):
                agregar(linea)

            try:
                tablas = page.extract_tables() or []
            except Exception:
                tablas = []

            for tabla in tablas:
                for fila in tabla or []:
                    celdas = [
                        re.sub(
                            r"\s+",
                            " ",
                            (celda or "").replace("\n", " "),
                        ).strip()
                        for celda in (fila or [])
                    ]
                    if any(celdas):
                        agregar(" | ".join(celdas), permitir_repetido=True)

            # El OCR se usa como respaldo para imágenes o texto mal codificado.
            texto_existente = f"{texto_layout}\n{texto_normal}".strip()
            requiere_ocr = (
                numero_pagina <= 2
                or len(re.sub(r"\s+", "", texto_existente)) < 120
            )
            if requiere_ocr:
                texto_ocr = _ocr_pagina_si_disponible(page)
                if texto_ocr:
                    agregar(texto_ocr)

    return "\n".join(lineas_salida)

# --- ANALIZADOR INTELIGENTE MULTILÍNEA GENERALIZADO ---
def analizar_pdf_inteligente(texto, metadatos_pdf=None):
    datos = {
        "nombre": "", "cargo": "", "tipo_examen": "PERIODICO",
        "examenes_lista": [], "recomendaciones_lista": [], "vigilancia_lista": [],
        "observaciones": "", "remisiones": "No", "consecutivo": "",
        "vigilancia_programa": "NINGUNO",
        "lugar": "",
        "fecha": datetime.date.today()
    }
    if not texto:
        return datos

    lineas_raw = texto.split("\n")

    # Extracción prioritaria y robusta de los tres campos que suelen venir
    # dentro de tablas: trabajador, cargo y lugar de realización.
    identificacion = extraer_identidad_cargo_lugar(texto)

    if identificacion.get("nombre"):
        datos["nombre"] = identificacion["nombre"]

    if identificacion.get("cargo"):
        datos["cargo"] = identificacion["cargo"]

    if identificacion.get("lugar"):
        datos["lugar"] = identificacion["lugar"]

    if identificacion.get("fecha"):
        datos["fecha"] = identificacion["fecha"]

    # Los valores obtenidos directamente desde celdas y coordenadas del PDF
    # tienen prioridad sobre el texto plano, porque conservan mejor la tabla.
    metadatos_pdf = metadatos_pdf or {}
    if metadatos_pdf.get("nombre"):
        datos["nombre"] = metadatos_pdf["nombre"]
    if metadatos_pdf.get("cargo"):
        datos["cargo"] = metadatos_pdf["cargo"]
    if metadatos_pdf.get("lugar"):
        datos["lugar"] = metadatos_pdf["lugar"]
    if metadatos_pdf.get("fecha"):
        datos["fecha"] = metadatos_pdf["fecha"]

    EXAMS_MAP = {
        "AUDIOMETRIA DE TONOS": "Audiometría", "AUDIOMETRIA": "Audiometría",
        "ESPIROMETRIA": "Espirometría", "ESPIROMETRÍA": "Espirometría",
        "OPTOMETRIA": "Optometría", "OPTOMETRÍA": "Optometría",
        "VISIOMETRIA": "Visiometría", "VISIOMETRÍA": "Visiometría",
        "EXAMEN MEDICO OCUPACIONAL": "Examen Clínico Ocupacional",
        "EXAMEN MEDICO": "Examen Clínico Ocupacional",
        "EXAMEN OCUPACIONAL ENFASIS OSTEOMUSCULAR": "Énfasis Osteomuscular",
        "ENFASIS OSTEOMUSCULAR": "Énfasis Osteomuscular", "ÉNFASIS OSTEOMUSCULAR": "Énfasis Osteomuscular",
        "ELECTROCARDIOGRAMA DE RITMO": "Electrocardiograma", "ELECTROCARDIOGRAMA": "Electrocardiograma", 
        "FROTIS": "Frotis",
        "CUADRO HEMATICO": "Cuadro Hemático", "CUADRO HEMÁTICO": "Cuadro Hemático",
        "COLESTEROL": "Colesterol",
        "TRIGLICERIDOS": "Triglicéridos", "TRIGLICÉRIDOS": "Triglicéridos",
        "PARCIAL DE ORINA": "Parcial de Orina",
        "VSH": "VSH", "PCR": "PCR"
    }

    examenes_detectados = []
    recoms_raw_dict = {}
    current_exam = None
    in_exams_section = True
    formato_grilla_detectado = False
    recoms_grilla_acumuladas = []

    for idx_l, linea in enumerate(lineas_raw):
        linea_limpia = limpiar_linea_ruido_lateral(linea)
        linea_upper = linea_limpia.upper().strip()
        
        if "EL CONCEPTO DE APTITUD SE DEFINIÓ A PARTIR DE LOS SIGUIENTES EXÁMENES PRACTICADOS" in linea_upper:
            for offset in range(1, 4):
                if idx_l + offset < len(lineas_raw):
                    l_sig = lineas_raw[idx_l + offset].upper()
                    for k_ex, v_ex in EXAMS_MAP.items():
                        if k_ex in l_sig and v_ex not in examenes_detectados:
                            examenes_detectados.append(v_ex)
            continue
            
        if any(h in linea_upper for h in ["RECOMENDACIONES MÉDICAS", "RECOMENDACIONES OCUPACIONALES", "HABITOS Y ESTILO DE VIDA SALUDABLES"]):
            formato_grilla_detectado = True
            continue
            
        if formato_grilla_detectado:
            if (
                any(
                    stop in linea_upper
                    for stop in [
                        "OTRAS OBSERVACIONES",
                        "REMISIONES:",
                        "ATENTAMENTE",
                        "CONSENTIMIENTO",
                        "AUTORIZO",
                        "TRATAMIENTO DE DATOS",
                        "HABEAS DATA",
                        "FIRMA DEL TRABAJADOR",
                        "FIRMA DEL PACIENTE",
                        "DECLARO",
                        "MANIFIESTO",
                    ]
                )
                or es_encabezado_legal(linea_limpia)
                or es_contenido_legal_recomendacion(linea_limpia)
            ):
                formato_grilla_detectado = False
                continue
            else:
                columnas = [
                    col.strip(" |/-,_.")
                    for col in re.split(r'\s{2,}|\|', linea_limpia)
                    if col.strip()
                ]
                for col in columnas:
                    col_clinica = recortar_contenido_legal(col)
                    if (
                        col_clinica
                        and not es_vacio_o_estado(col_clinica)
                        and not es_contenido_legal_recomendacion(col_clinica)
                    ):
                        rec_fmt = a_caso_oracion(col_clinica)
                        if rec_fmt and rec_fmt not in recoms_grilla_acumuladas:
                            recoms_grilla_acumuladas.append(rec_fmt)
                continue

        if (
            any(
                stop in linea_upper
                for stop in [
                    "OBSERVACIONES:",
                    "OBSERVACION:",
                    "REMISIONES:",
                    "SISTEMA DE VIGILANCIA",
                    "CONSENTIMIENTO",
                    "AUTORIZO",
                    "TRATAMIENTO DE DATOS",
                    "HABEAS DATA",
                    "FIRMA DEL TRABAJADOR",
                    "FIRMA DEL PACIENTE",
                    "DECLARO",
                    "MANIFIESTO",
                    "ATENTAMENTE",
                ]
            )
            or es_encabezado_legal(linea_limpia)
            or es_contenido_legal_recomendacion(linea_limpia)
        ):
            in_exams_section = False
            if current_exam:
                contenido_actual = recoms_raw_dict.get(current_exam, "")
                recoms_raw_dict[current_exam] = recortar_contenido_legal(
                    contenido_actual
                )
                current_exam = None
            continue

        matched_key = None
        for key in sorted(EXAMS_MAP.keys(), key=len, reverse=True):
            if key in linea_upper and linea_upper.find(key) < 15:
                matched_key = key
                break
        
        if matched_key:
            in_exams_section = True
            current_exam = EXAMS_MAP[matched_key]
            if current_exam not in examenes_detectados:
                examenes_detectados.append(current_exam)
            idx = linea_upper.find(matched_key) + len(matched_key)
            recoms_raw_dict[current_exam] = linea_limpia[idx:].strip(" :-,_/")
        else:
            if in_exams_section and current_exam and linea_limpia.strip():
                if (
                    not es_encabezado_legal(linea_limpia)
                    and not es_contenido_legal_recomendacion(linea_limpia)
                    and not (linea_limpia.isupper() and len(linea_limpia) > 10)
                ):
                    recoms_raw_dict[current_exam] = (
                        recoms_raw_dict.get(current_exam, "")
                        + " "
                        + linea_limpia.strip()
                    )

    recoms_por_examen = []
    pve_detectados = set()

    if recoms_grilla_acumuladas:
        for rec in recoms_grilla_acumuladas:
            recoms_por_examen.append(rec)
            rec_up = rec.upper()
            if any(re.search(patron, rec_up) for patron in [r'\bAUDITIV', r'\bRUIDO', r'\bOIDO', r'\bOÍDO', r'\bAUDIO']): pve_detectados.add("Conservación Auditiva")
            if any(re.search(patron, rec_up) for patron in [r'\bPOSTURAL', r'\bLUMBAR', r'\bOSTEOMUSCULAR', r'\bERGONOMIC', r'\bESPALDA', r'\bCARGA']): pve_detectados.add("Prevención Osteomuscular (DME)")
            if any(re.search(patron, rec_up) for patron in [r'\bVISUAL', r'\bGAFAS', r'\bVISION', r'\bVISIÓN', r'\bLENTE', r'\bOPTOMETR', r'\bRX\b']): pve_detectados.add("Conservación Visual")
            if any(re.search(patron, rec_up) for patron in [r'\bRESPIRATORI', r'\bESPIROMETR', r'\bPOLVO', r'\bHUMO']): pve_detectados.add("Conservación Respiratoria")
    else:
        for exam in examenes_detectados:
            rec_part = recoms_raw_dict.get(exam, "").strip()
            rec_part = recortar_contenido_legal(rec_part)
            rec_part = re.sub(r'\s+', ' ', rec_part)
            rec_part = limpiar_ruido_columnas_final(rec_part)
            
            if not es_vacio_o_estado(rec_part):
                parts = re.split(r'//|;|\b\d+\.|\b\d+\-', rec_part)
                valid_parts = []
                for p in parts:
                    p_clean = p.strip(" .-_/()[]")
                    p_clean = recortar_contenido_legal(p_clean)
                    if (
                        p_clean
                        and not es_vacio_o_estado(p_clean)
                        and not es_contenido_legal_recomendacion(p_clean)
                    ):
                        valid_parts.append(a_caso_oracion(p_clean))
                        p_upper = p_clean.upper()
                        if any(re.search(patron, p_upper) for patron in [r'\bAUDITIV', r'\bRUIDO', r'\bOIDO', r'\bOÍDO', r'\bAUDIO']): pve_detectados.add("Conservación Auditiva")
                        if any(re.search(patron, p_upper) for patron in [r'\bPOSTURAL', r'\bLUMBAR', r'\bOSTEOMUSCULAR', r'\bERGONOMIC', r'\bESPALDA', r'\bCARGA']): pve_detectados.add("Prevención Osteomuscular (DME)")
                        if any(re.search(patron, p_upper) for patron in [r'\bVISUAL', r'\bGAFAS', r'\bVISION', r'\bVISIÓN', r'\bLENTE', r'\bOPTOMETR', r'\bRX\b']): pve_detectados.add("Conservación Visual")
                        if any(re.search(patron, p_upper) for patron in [r'\bRESPIRATORI', r'\bESPIROMETR', r'\bPOLVO', r'\bHUMO']): pve_detectados.add("Conservación Respiratoria")
                
                if valid_parts:
                    recoms_por_examen.append(f"{exam}: {' - '.join(valid_parts)}")

    # CLAVE UNIFICADA EN ESPAÑOL DEFINITIVA
    datos["examenes_lista"] = examenes_detectados
    datos["recomendaciones_lista"] = filtrar_recomendaciones_clinicas(
        recoms_por_examen
    )
    datos["vigilancia_lista"] = list(pve_detectados)

    # --- CORRECCIÓN INTEGRAL: EXTRACCIÓN ACOTA DE PROGRAMAS PVE ---
    programas_encontrados = []
    sve_clinical_keywords = {
        "VISUAL": "Conservación Visual", "AUDITIV": "Conservación Auditiva", 
        "RUIDO": "Conservación Auditiva", "OIDO": "Conservación Auditiva", "OÍDO": "Conservación Auditiva",
        "AUDIO": "Conservación Auditiva", "OSTEOMUSCULAR": "Prevención Osteomuscular (DME)",
        "POSTURAL": "Prevención Osteomuscular (DME)", "LUMBAR": "Prevención Osteomuscular (DME)",
        "ERGONOMIC": "Prevención Osteomuscular (DME)", "ESPALDA": "Prevención Osteomuscular (DME)",
        "DME": "Prevención Osteomuscular (DME)", "RESPIRATORI": "Conservación Respiratoria",
        "ESPIROMETR": "Conservación Respiratoria", "POLVO": "Conservación Respiratoria",
        "HUMO": "Conservación Respiratoria", "CARDIOVASCULAR": "Riesgo Cardiovascular"
    }

    # Búsqueda vecinal: Solo lee las 2 líneas inmediatamente inferiores a la cabecera PVE
    for idx, line in enumerate(lineas_raw):
        l_up = line.upper()
        if "INGRESAR AL PROGRAMA DE VIGILANCIA" in l_up or "PROGRAMA DE VIGILANCIA EPIDEMIOL" in l_up:
            for offset in [0, 1, 2]:
                if idx + offset < len(lineas_raw):
                    text_target = lineas_raw[idx + offset].upper()
                    # Freno de emergencia si cruza a la declaración de firmas/consentimientos legales
                    if offset > 0 and any(stop in text_target for stop in ["REMISIONES:", "OBSERVACIONES:", "ATENTAMENTE", "CONSENTIMIENTO"]):
                        break
                    for kw, prog_name in sve_clinical_keywords.items():
                        if kw in text_target and prog_name not in programas_encontrados:
                            programas_encontrados.append(prog_name)
            break

    if not programas_encontrados and pve_detectados:
        programas_encontrados = [p.upper() for p in pve_detectados]
        
    datos["vigilancia_programa"] = ", ".join(programas_encontrados) if programas_encontrados else "NINGUNO"

    def extraer_seccion_limpia(texto_completo, palabras_inicio, palabras_fin):
        seccion = []
        dentro = False
        for l in texto_completo.split('\n'):
            l_limpia = limpiar_linea_ruido_lateral(l)
            l_upper = l_limpia.upper().strip()
            if not dentro:
                if any(h in l_upper for h in palabras_inicio):
                    dentro = True
                    for h in palabras_inicio:
                        if h in l_upper:
                            resto = l_limpia[l_upper.find(h) + len(h):].strip(" :-,_")
                            if resto: seccion.append(resto)
                            break
            else:
                if any(h in l_upper for h in palabras_fin): break
                seccion.append(l_limpia)
        return "\n".join([s for s in seccion if s]).strip()

    obs_fmt_nuevo = ""
    m_obs_nuevo = re.search(r'OTRAS OBSERVACIONES Y RECOMENDACIONES\s*\n\s*([^\n]+)', texto, re.IGNORECASE)
    if m_obs_nuevo: obs_fmt_nuevo = m_obs_nuevo.group(1).strip()
        
    if obs_fmt_nuevo and not es_vacio_o_estado(obs_fmt_nuevo):
        datos["observaciones"] = a_caso_oracion(obs_fmt_nuevo)
    else:
        datos["observaciones"] = a_caso_oracion(extraer_seccion_limpia(
            texto, ["OBSERVACIONES:"], ["RECOMENDACIONES", "REMISIONES", "INGRESAR AL PROGRAMA", "PROGRAMA DE VIGILANCIA"]
        ))
    
    rem_raw = extraer_seccion_limpia(texto, ["INFORMACION DE REMISIONES", "INFORMACIÓN DE REMISIONES"], ["CONSENTIMIENTO", "AUTORIZO"])
    datos["remisiones"] = "No" if es_vacio_o_negativo(rem_raw) else a_caso_oracion(rem_raw)

    return datos

# --- FORMATEADOR DINÁMICO DE NEGRITAS EN EL CUERPO ---
def aplicar_negrita_dinamica_cuerpo(paragraph, tipo_examen):
    texto_parrafo = paragraph.text
    if "Según los lineamientos del programa de medicina preventiva" not in texto_parrafo:
        return
        
    paragraph.text = "" 
    p1 = "Según los lineamientos del programa de medicina preventiva y del trabajo de JER S.A; se hace entrega de las recomendaciones establecidas por el Proveedor de servicios de Exámenes Médico Ocupacionales ("
    paragraph.add_run(p1)
    
    opciones = [
        ("Ingreso", "INGRESO" in tipo_examen.upper()),
        ("Periódico", "PERIODIC" in tipo_examen.upper() or "PERIÓDIC" in tipo_examen.upper()),
        ("egreso", "EGRESO" in tipo_examen.upper() or "RETIRO" in tipo_examen.upper()),
        ("cambio de cargo", "CAMBIO" in tipo_examen.upper()),
        ("post incapacidad", "POST" in tipo_examen.upper() or "INCAPACIDAD" in tipo_examen.upper())
    ]
    
    for i, (texto_opcion, condicion) in enumerate(opciones):
        run = paragraph.add_run(texto_opcion)
        run.bold = condicion 
        if i < len(opciones) - 1:
            if i == len(opciones) - 2: paragraph.add_run(" y ")
            else: paragraph.add_run(", ")
                
    paragraph.add_run(")")

def replace_placeholder_in_paragraph_runs(paragraph, placeholder, value):
    if placeholder not in paragraph.text: return False
    replaced = False
    for run in paragraph.runs:
        if placeholder in run.text:
            run.text = run.text.replace(placeholder, value)
            replaced = True
            
    if not replaced:
        font_name = "Arial"; font_size = Pt(11); bold = False; italic = False; color = None
        if paragraph.runs:
            for r in paragraph.runs:
                if r.text.strip():
                    font_name = r.font.name or font_name
                    font_size = r.font.size or font_size
                    bold = r.bold if r.bold is not None else bold
                    italic = r.italic if r.italic is not None else italic
                    color = r.font.color.rgb if r.font.color else color
                    break
        full_text = paragraph.text.replace(placeholder, value)
        paragraph.text = ""
        new_run = paragraph.add_run(full_text)
        new_run.font.name = font_name; new_run.font.size = font_size; new_run.bold = bold; new_run.italic = italic
        if color: new_run.font.color.rgb = color
    return True

def insert_bullets_in_placeholder(parent_container, paragraph, items_list):
    if not paragraph.runs:
        font_name = "Arial"; font_size = Pt(11); bold = False; color = None
    else:
        first_run = paragraph.runs[0]
        font_name = first_run.font.name or "Arial"
        font_size = first_run.font.size or Pt(11)
        bold = first_run.bold
        color = first_run.font.color.rgb if first_run.font.color else None

    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.space_before = Pt(0)
    
    if not items_list:
        run = paragraph.add_run("Ninguno.")
        run.font.name = font_name; run.font.size = font_size; run.bold = bold
        if color: run.font.color.rgb = color
        return

    run = paragraph.add_run("• " + items_list[0])
    run.font.name = font_name; run.font.size = font_size; run.bold = bold
    if color: run.font.color.rgb = color

    current_p = paragraph
    for item in items_list[1:]:
        new_p_element = OxmlElement('w:p')
        current_p._p.addnext(new_p_element)
        new_para = Paragraph(new_p_element, parent_container)
        new_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
        new_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
        new_para.paragraph_format.space_after = Pt(2)
        new_para.paragraph_format.space_before = Pt(0)
        new_para.paragraph_format.left_indent = paragraph.paragraph_format.left_indent or Inches(0.25)
        run_new = new_para.add_run("• " + item)
        run_new.font.name = font_name; run_new.font.size = font_size; run_new.bold = bold
        if color: run_new.font.color.rgb = color
        current_p = new_para

def insert_recommendations_in_placeholder(parent_container, paragraph, recom_list):
    combined_items = list(recom_list)
    if paragraph.runs:
        font_name = paragraph.runs[0].font.name or "Arial"
        font_size = paragraph.runs[0].font.size or Pt(11)
        color = paragraph.runs[0].font.color.rgb if paragraph.runs[0].font.color else None
    else:
        font_name = "Arial"; font_size = Pt(11); color = None

    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.space_before = Pt(0)
    run_lbl = paragraph.add_run("Recomendaciones: ")
    run_lbl.bold = True; run_lbl.font.name = font_name; run_lbl.font.size = font_size
    if color: run_lbl.font.color.rgb = color

    if not combined_items:
        run_none = paragraph.add_run("Ninguna.")
        run_none.font.name = font_name; run_none.font.size = font_size
        if color: run_none.font.color.rgb = color
        return

    run_first = paragraph.add_run("• " + combined_items[0])
    run_first.font.name = font_name; run_first.font.size = font_size
    if color: run_first.font.color.rgb = color

    current_p = paragraph
    for item in combined_items[1:]:
        new_p_element = OxmlElement('w:p')
        current_p._p.addnext(new_p_element)
        new_para = Paragraph(new_p_element, parent_container)
        new_para.paragraph_format.alignment = paragraph.paragraph_format.alignment
        new_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing
        new_para.paragraph_format.space_after = Pt(2)
        new_para.paragraph_format.space_before = Pt(0)
        new_para.paragraph_format.left_indent = Inches(0.25)
        run_new = new_para.add_run("• " + item)
        run_new.font.name = font_name; run_new.font.size = font_size
        if color: run_new.font.color.rgb = color
        current_p = new_para

def replace_label_placeholder(paragraph, label_text, placeholder, value):
    if placeholder not in paragraph.text: return False
    if paragraph.runs:
        font_name = paragraph.runs[0].font.name or "Arial"
        font_size = paragraph.runs[0].font.size or Pt(11)
        color = paragraph.runs[0].font.color.rgb if paragraph.runs[0].font.color else None
    else:
        font_name = "Arial"; font_size = Pt(11); color = None
        
    paragraph.text = ""
    paragraph.paragraph_format.space_after = Pt(3)
    paragraph.paragraph_format.space_before = Pt(0)
    run_lbl = paragraph.add_run(label_text)
    run_lbl.bold = True; run_lbl.font.name = font_name; run_lbl.font.size = font_size
    if color: run_lbl.font.color.rgb = color
    
    val_clean = value.strip() if value else "Ninguna."
    if es_vacio_o_negativo(val_clean): val_clean = "Ninguna."
        
    run_val = paragraph.add_run(val_clean)
    run_val.font.name = font_name; run_val.font.size = font_size
    if color: run_val.font.color.rgb = color
    return True

def obtener_siguiente_consecutivo_local():
    val = obtener_config("ultimo_consecutivo_local")
    return int(val) + 1 if val else 1

def incrementar_consecutivo_local():
    next_num = obtener_siguiente_consecutivo_local()
    guardar_config("ultimo_consecutivo_local", str(next_num))
    return f"SST-2026-{next_num}"

# --- CONSTRUCTOR DE DOCUMENTO ÚNICO INTELIGENTE ---
def generar_word_unico(datos_trabajador, lugar, fecha, template_uploaded, firma_file):
    if template_uploaded: doc_word = Document(template_uploaded)
    elif os.path.exists("FORMATO RECOMENDACIONES MEDICAS BOT.docx"): doc_word = Document("FORMATO RECOMENDACIONES MEDICAS BOT.docx")
    else: doc_word = Document()
    
    consecutivo_final = datos_trabajador.get("consecutivo", "")
    if not consecutivo_final:
        g_url = obtener_config("google_sheets_url")
        if g_url:
            try:
                r = requests.get(g_url, params={
                    "name": datos_trabajador["nombre"], "cargo": datos_trabajador["cargo"], 
                    "examen": datos_trabajador["tipo_examen"], "fecha": fecha.strftime("%Y-%m-%d")
                }, timeout=10)
                consecutivo_final = r.json().get("consecutive") if r.json().get("status") == "success" else incrementar_consecutivo_local()
            except: consecutivo_final = incrementar_consecutivo_local()
        else: consecutivo_final = incrementar_consecutivo_local()
        datos_trabajador["consecutivo"] = consecutivo_final

    simple_replacements = {
        "{{NUMERO DE CONSECUTIVO}}": consecutivo_final, 
        "{{TIPO DE EXAMEN}}": datos_trabajador["tipo_examen"].upper(),
        "{{LUGAR}}": lugar, "{{FECHA HOY}}": fecha.strftime("%d de %B de %Y"),
        "{{NOMBRE DE LA PERSONA}}": datos_trabajador["nombre"].upper(), 
        "{{CARGO DE LA PERSONA}}": datos_trabajador["cargo"].upper(),
        "{{Programa de vigilancia epidemiológica}}": datos_trabajador.get("vigilancia_programa", "NINGUNO").upper()
    }

    def procesar_parrafo(p, container):
        if "{{LISTA DE EXAMENES REALIZADOS}}" in p.text:
            insert_bullets_in_placeholder(container, p, datos_trabajador["examenes_lista"])
            return True
        if "{{Recomendaciones médicas}}" in p.text:
            insert_recommendations_in_placeholder(container, p, datos_trabajador["recomendaciones_lista"])
            return True
        if "{{Observaciones}}".lower() in p.text.lower():
            replace_label_placeholder(p, "Observaciones: ", p.text, datos_trabajador["observaciones"])
            return True
        if "{{Remisiones}}".lower() in p.text.lower():
            replace_label_placeholder(p, "Remisiones: ", p.text, datos_trabajador["remisiones"])
            return True
            
        for k, v in simple_replacements.items():
            if k in p.text: replace_placeholder_in_paragraph_runs(p, k, v)
        aplicar_negrita_dinamica_cuerpo(p, datos_trabajador["tipo_examen"])

    for p in list(doc_word.paragraphs): procesar_parrafo(p, doc_word)

    for table in doc_word.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in list(cell.paragraphs): procesar_parrafo(p, cell)
                
                idx_victor = -1
                for idx, p in enumerate(cell.paragraphs):
                    if "VÍCTOR ALONSO MORENO CASAS" in p.text:
                        idx_victor = idx
                        break
                if idx_victor != -1 and firma_file:
                    target_idx = max(0, idx_victor - 1)
                    p_firma = cell.paragraphs[target_idx]
                    p_firma.text = ""
                    p_firma.add_run().add_picture(firma_file, width=Inches(1.6))

    b_io = io.BytesIO()
    doc_word.save(b_io)
    return b_io.getvalue(), consecutivo_final

# --- COMPILADOR DE WORD A PDF ---
def convertir_docx_a_pdf(docx_bytes):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as temp_docx:
        temp_docx.write(docx_bytes)
        temp_docx_path = temp_docx.name
    pdf_path = temp_docx_path.replace(".docx", ".pdf")
    
    try:
        subprocess.run([
            "libreoffice", "--headless", "--convert-to", "pdf", 
            "--outdir", os.path.dirname(temp_docx_path), temp_docx_path
        ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL, timeout=15)
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f: pdf_bytes = f.read()
            os.unlink(temp_docx_path); os.unlink(pdf_path)
            return pdf_bytes, True
    except: pass

    try:
        from docx2pdf import convert
        convert(temp_docx_path, pdf_path)
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f: pdf_bytes = f.read()
            os.unlink(temp_docx_path); os.unlink(pdf_path)
            return pdf_bytes, True
    except: pass

    if os.path.exists(temp_docx_path): os.unlink(temp_docx_path)
    return None, False

def generar_html_vista(datos, consecutivo_num, lugar, fecha):
    return f"""
    <div style="font-family: Arial, sans-serif; color: #333; padding: 20px; line-height: 1.5; background: white; border: 1px solid #ccc; max-width: 800px; margin: auto;">
        <div style="text-align: right; font-weight: bold; color: #1f4e79;">Consecutivo: {consecutivo_num}</div>
        <div style="text-align: center; font-weight: bold; font-size: 16px; margin: 20px 0; color: #1f4e79; background: #f0f4f8; padding: 8px;">
            ASUNTO: RECOMENDACIONES EXAMEN {datos['tipo_examen'].upper()}
        </div>
        <div>{lugar}, {fecha.strftime('%d de %B de %Y')}</div><br>
        <div>Sr(a).<br><strong>{datos['nombre']}</strong><br>{datos['cargo']}</div><br>
        <p>Cordial saludo,</p>
        <p>Según los lineamientos del programa de medicina preventiva y del trabajo de JER S.A; se hace entrega de las recomendaciones establecidas por el Proveedor de servicios de Exámenes Médico Ocupacionales (Ingreso, Periódico, egreso, cambio de cargo y post incapacidad)</p>
        <p><strong>EXÁMENES REALIZADOS:</strong></p>
        <ul>{"".join([f"<li>{ex}</li>" for ex in datos['examenes_lista']])}</ul>
        <p><strong>Recomendaciones:</strong></p>
        <ul>{"".join([f"<li>{rec}</li>" for rec in datos['recomendaciones_lista']])}</ul>
        <p><strong>Programa de Vigilancia:</strong> {datos.get('vigilancia_programa', 'NINGUNO')}</p>
        <p><strong>observaciones:</strong> {datos['observaciones']}</p>
        <p><strong>remisiones:</strong> {datos['remisiones']}</p><br>
        <p>Atentamente,</p><br>
        <p><strong>VÍCTOR ALONSO MORENO CASAS</strong><br>Coordinador SST</p>
    </div>
    """

# --- VISTA PRINCIPAL STREAMLIT ---
st.markdown("<div class='header-banner'><h1>🩺 Portal de Control SST - JER S.A.</h1><p>Generación de Comunicaciones con Negrita Dinámica, Google Sheets y Firma Digital</p></div>", unsafe_allow_html=True)

st.sidebar.markdown(f"<h3 style='color:#60a5fa;'>👤 Perfil Activo</h3>", unsafe_allow_html=True)
st.sidebar.markdown(f"<div class='metric-card'><strong>Usuario:</strong> {st.session_state.username}</div>", unsafe_allow_html=True)
if st.sidebar.button("Cerrar Sesión"):
    st.session_state.logged_in = False
    st.session_state.username = ""
    st.session_state.documentos = {}
    st.session_state.pdfs_raw_bytes = {}
    st.session_state.textos_raw = {}
    st.session_state.export_bytes = None
    st.session_state.zip_bytes = None
    st.session_state.processed_doc = None
    st.session_state.prev_colaborador = None
    st.session_state.document_count = 0
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.subheader("🔗 Configuración de Nube Sheets")
g_url_guardada = obtener_config("google_sheets_url")
g_url_input = st.sidebar.text_input("URL de Google Apps Script:", value=g_url_guardada, type="password")
if st.sidebar.button("Guardar Conexión"):
    guardar_config("google_sheets_url", g_url_input)
    st.sidebar.success("¡Enlace guardado!")
    st.rerun()

st.sidebar.markdown("---")
st.sidebar.subheader("📋 Documentación Base")
template_uploaded = st.sidebar.file_uploader("Formato de Word Institucional (.docx)", type=["docx"])
firma_file = st.sidebar.file_uploader("Estampa de Firma Autorizada (.png / .jpg)", type=["png", "jpg"])

col_izq, col_der = st.columns([1, 1.2])

with col_izq:
    st.markdown("<h3 style='color:#60a5fa;'>📂 1. Carga de Documentos PDF</h3>", unsafe_allow_html=True)
    pdfs_subidos = st.file_uploader("Carga los archivos PDF:", type="pdf", accept_multiple_files=True)
    
    if pdfs_subidos:
        if len(pdfs_subidos) != st.session_state.document_count:
            st.session_state.documentos = {}
            st.session_state.pdfs_raw_bytes = {}
            st.session_state.zip_bytes = None
            st.session_state.document_count = len(pdfs_subidos)
            
        for pdf in pdfs_subidos:
            if pdf.name not in st.session_state.documentos:
                pdf_raw_data = pdf.read()
                st.session_state.pdfs_raw_bytes[pdf.name] = pdf_raw_data
                texto_raw = extraer_texto_pdf_robusto(pdf_raw_data)
                metadatos_pdf = extraer_metadatos_pdf_estructurados(
                    pdf_raw_data,
                    texto_raw,
                )
                st.session_state.textos_raw[pdf.name] = texto_raw
                st.session_state.documentos[pdf.name] = analizar_pdf_inteligente(
                    texto_raw,
                    metadatos_pdf=metadatos_pdf,
                )
        
        st.markdown(f"""
            <div style='display:flex; gap:10px; margin-top:15px;'>
                <div class='metric-card' style='flex:1;'><strong>PDFs Leídos</strong><br><span style='font-size:20px; font-weight:700; color:#60a5fa;'>{len(st.session_state.documentos)}</span></div>
                <div class='metric-card' style='flex:1;'><strong>Nube Sheets</strong><br><span style='font-size:14px; font-weight:600; color:#4ade80;'>{'Conectado' if g_url_guardada else 'Modo Local'}</span></div>
            </div>
        """, unsafe_allow_html=True)
        
        archivo_seleccionado = st.selectbox("🎯 Selecciona Colaborador:", list(st.session_state.documentos.keys()))
        
        if archivo_seleccionado and archivo_seleccionado in st.session_state.pdfs_raw_bytes:
            st.markdown("---")
            st.markdown("<h4 style='color:#60a5fa;'>📄 Soporte Visual de Comparación</h4>", unsafe_allow_html=True)
            
            bytes_originales = st.session_state.pdfs_raw_bytes[archivo_seleccionado]
            st.download_button(
                label="📥 Descargar PDF de Origen (IPS)",
                data=bytes_originales,
                file_name=f"ORIGINAL_{archivo_seleccionado}",
                mime="application/pdf",
                key="btn_download_original_source"
            )
            
            with st.expander("👁️ Ver / Ocultar PDF de Origen Subido", expanded=True):
                with pdfplumber.open(io.BytesIO(bytes_originales)) as preview_pdf:
                    for i, page in enumerate(preview_pdf.pages):
                        img_pil = page.to_image(resolution=130).original
                        st.image(img_pil, caption=f"Página {i+1} - {archivo_seleccionado}", use_container_width=True)
    else:
        archivo_seleccionado = None
        st.session_state.documentos = {}
        st.session_state.pdfs_raw_bytes = {}
        st.session_state.zip_bytes = None
        st.session_state.document_count = 0

with col_der:
    st.markdown("<h3 style='color:#60a5fa;'>📋 2. Editor del Trabajador Seleccionado</h3>", unsafe_allow_html=True)
    if archivo_seleccionado:
        if archivo_seleccionado != st.session_state.prev_colaborador:
            st.session_state.processed_doc = None
            st.session_state.prev_colaborador = archivo_seleccionado
            
        doc_actual = st.session_state.documentos[archivo_seleccionado]
        
        col_f1, col_f2 = st.columns(2)
        st.caption("Verifica nombre, cargo, fecha y lugar antes de generar el documento. Los campos siguen siendo editables.")
        with col_f1: lugar = st.text_input("Lugar:", value=doc_actual.get("lugar", ""))
        with col_f2: fecha = st.date_input("Fecha:", value=doc_actual.get("fecha", datetime.date.today()))
        
        tipo_examen = st.text_input("Tipo de Examen:", value=doc_actual["tipo_examen"].upper())
        
        col_p1, col_p2 = st.columns(2)
        with col_p1: nombre_persona = st.text_input("Trabajador:", value=doc_actual["nombre"])
        with col_p2: cargo_persona = st.text_input("Cargo:", value=doc_actual["cargo"])
        
        examenes_realizados = st.text_area("Exámenes Realizados:", value="\n".join(doc_actual["examenes_lista"]))
        recom_medicas = st.text_area("Recomendaciones por Examen:", value="\n".join(doc_actual["recomendaciones_lista"]), height=130)
        
        programa_vigilancia = st.text_input("Programa de Vigilancia Epidemiológica (PVE):", value=doc_actual.get("vigilancia_programa", "NINGUNO"))
        
        observaciones = st.text_area("Observaciones:", value=doc_actual["observaciones"])
        remisiones = st.text_input("Remisiones (Escribe 'No' para marcarlo negativo):", value=doc_actual["remisiones"])

        valores_actualizados = {
            "nombre": nombre_persona, "cargo": cargo_persona, "tipo_examen": tipo_examen,
            "examenes_lista": [l.strip() for l in examenes_realizados.split('\n') if l.strip()],
            "recomendaciones_lista": [l.strip() for l in recom_medicas.split('\n') if l.strip()],
            "observaciones": observaciones.strip(), "remisiones": remisiones.strip(),
            "vigilancia_programa": programa_vigilancia.strip(),
            "lugar": lugar, "fecha": fecha
        }
        
        for clave, valor in valores_actualizados.items():
            if doc_actual.get(clave) != valor:
                st.session_state.processed_doc = None
                doc_actual[clave] = valor

        st.markdown("---")
        formato_salida = st.radio("⚡ Elige formato de generación:", ["Microsoft Word (.docx)", "Documento PDF Oficial (.pdf)", "Impresión de Respaldo Web (HTML)"], horizontal=True)
        
        col_act1, col_gen2 = st.columns(2)
        
        with col_act1:
            if st.button("✨ Procesar este Colaborador"):
                with st.spinner("Procesando y registrando documento..."):
                    bytes_word, consec_num = generar_word_unico(doc_actual, lugar, fecha, template_uploaded, firma_file)
                    
                    if "Word" in formato_salida:
                        st.session_state.processed_doc = {
                            "bytes": bytes_word, "consec_num": consec_num,
                            "filename": f"Informe_{nombre_persona.replace(' ','_')}.docx",
                            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        }
                    elif "PDF" in formato_salida:
                        bytes_pdf, exito = convertir_docx_a_pdf(bytes_word)
                        if exito:
                            st.session_state.processed_doc = {
                                "bytes": bytes_pdf, "consec_num": consec_num,
                                "filename": f"Informe_{nombre_persona.replace(' ','_')}.pdf", "mime": "application/pdf"
                            }
                        else: st.error("⚠️ No se pudo compilar el PDF de manera directa para coincidir con el Word.")
                    else:
                        html_out = generar_html_vista(doc_actual, consec_num, lugar, fecha)
                        st.session_state.processed_doc = {
                            "bytes": html_out.encode('utf-8'), "consec_num": consec_num,
                            "filename": f"Informe_{nombre_persona.replace(' ','_')}.html", "mime": "text/html"
                        }
                st.rerun()
            
            if st.session_state.processed_doc is not None:
                doc_info = st.session_state.processed_doc
                st.success(f"🟢 Guardado con éxito en base (Consecutivo: {doc_info['consec_num']})")
                st.download_button(
                    label=f"📥 Descargar archivo generado ({doc_info['filename'].split('.')[-1].upper()})",
                    data=doc_info["bytes"], file_name=doc_info["filename"], mime=doc_info["mime"]
                )
                        
        with col_gen2:
            if len(st.session_state.documentos) > 1:
                if st.button("📦 Procesar TODOS los Colaboradores (ZIP)"):
                    with st.spinner("Compilando lote masivo..."):
                        zip_buffer = io.BytesIO()
                        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
                            for filename, datos_trab in st.session_state.documentos.items():
                                bytes_word, consec_num = generar_word_unico(datos_trab, lugar, fecha, template_uploaded, firma_file)
                                if "Word" in formato_salida:
                                    zf.writestr(f"Recomendaciones_{datos_trab['nombre'].replace(' ', '_')}.docx", bytes_word)
                                elif "PDF" in formato_salida:
                                    bytes_pdf, exito = convertir_docx_a_pdf(bytes_word)
                                    if exito: zf.writestr(f"Recomendaciones_{datos_trab['nombre'].replace(' ', '_')}.pdf", bytes_pdf)
                                    else: zf.writestr(f"Recomendaciones_{datos_trab['nombre'].replace(' ', '_')}.docx", bytes_word)
                                else:
                                    html_out = generar_html_vista(datos_trab, consec_num, lugar, fecha)
                                    zf.writestr(f"Recomendaciones_{datos_trab['nombre'].replace(' ', '_')}.html", html_out.encode('utf-8'))
                        zip_buffer.seek(0)
                        st.session_state.zip_bytes = zip_buffer.getvalue()
                        st.success("🎉 ZIP de lote masivo compilado con éxito.")
                        st.rerun()
                        
                if st.session_state.zip_bytes is not None:
                    st.download_button(
                        label="📥 Descargar ZIP Masivo", data=st.session_state.zip_bytes, 
                        file_name=f"Lote_SST_JER_SA_{fecha.strftime('%Y%m%d')}.zip", mime="application/zip"
                    )
    else:
        st.markdown("<div style='text-align:center; padding: 40px; color:#64748b;'><h3>👋 Tablero Listo</h3><p>Por favor, arrastra tus archivos PDF en la sección izquierda para activar el procesamiento automático.</p></div>", unsafe_allow_html=True)
